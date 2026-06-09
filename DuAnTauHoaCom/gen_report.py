# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

# ── Styles ────────────────────────────────────────────────────────
styles = doc.styles

def set_style(style_name, font_name='Times New Roman', size=12, bold=False, color=None):
    try:
        s = styles[style_name]
    except:
        return
    s.font.name = font_name
    s.font.size = Pt(size)
    s.font.bold = bold
    if color:
        s.font.color.rgb = RGBColor(*color)

set_style('Normal', size=12)
set_style('Heading 1', size=14, bold=True, color=(0,70,127))
set_style('Heading 2', size=13, bold=True, color=(0,112,192))
set_style('Heading 3', size=12, bold=True, color=(31,73,125))

# ── Helper functions ───────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def para(text, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(1 + level*0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.name = 'Times New Roman'
        c.paragraphs[0].runs[0].font.size = Pt(11)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'BDD7EE')
        c._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        r = t.add_row()
        for i, cell in enumerate(row):
            c = r.cells[i]
            c.text = str(cell)
            c.paragraphs[0].runs[0].font.name = 'Times New Roman'
            c.paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# TRANG BÌA
# ═══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('TRƯỜNG ĐẠI HỌC\nKHOA CÔNG NGHỆ THÔNG TIN')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('─────────────────────────────')
r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(30)
r3 = p3.add_run('BÁO CÁO PHÂN TÍCH DỰ ÁN')
r3.font.name = 'Times New Roman'; r3.font.size = Pt(16); r3.bold = True

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('HỆ THỐNG ĐẶT VÉ TÀU HỎA KLN TRAIN')
r4.font.name = 'Times New Roman'; r4.font.size = Pt(18); r4.bold = True
r4.font.color.rgb = RGBColor(0, 70, 127)

doc.add_paragraph()
doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run('Phân tích toàn diện kiến trúc, luồng chức năng\nvà cơ sở dữ liệu hệ thống đặt vé')
r5.font.name = 'Times New Roman'; r5.font.size = Pt(13); r5.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run('Năm học: 2025 – 2026\nNgày lập: 27/05/2026')
r6.font.name = 'Times New Roman'; r6.font.size = Pt(12)

page_break()

# ═══════════════════════════════════════════════════════════════════
# MỤC LỤC (thủ công)
# ═══════════════════════════════════════════════════════════════════
h1('MỤC LỤC')
toc_items = [
    ('I.', 'TỔNG QUAN DỰ ÁN', 3),
    ('II.', 'KIẾN TRÚC TỔNG THỂ', 4),
    ('III.', 'CƠ SỞ DỮ LIỆU – SQL SERVER KLN_TRAIN', 5),
    ('  3.1', 'Danh sách bảng và mục đích', 5),
    ('  3.2', 'Mô tả chi tiết từng bảng', 5),
    ('  3.3', 'Quan hệ giữa các bảng', 9),
    ('IV.', 'BACKEND – NODE.JS / EXPRESS', 10),
    ('  4.1', 'Điểm khởi động: app.js', 10),
    ('  4.2', 'Cấu hình CSDL: config/database.js', 10),
    ('  4.3', 'Models – Sequelize ORM', 11),
    ('  4.4', 'Repositories – Lớp truy cập dữ liệu', 12),
    ('  4.5', 'Services – Lớp nghiệp vụ', 14),
    ('  4.6', 'Controllers – Xử lý HTTP Request', 19),
    ('  4.7', 'Routes – Định tuyến API', 21),
    ('  4.8', 'Middleware', 22),
    ('  4.9', 'Utils – Tiện ích', 22),
    ('V.', 'FRONTEND – REACT 19 / VITE', 23),
    ('  5.1', 'Cấu trúc và điểm khởi động', 23),
    ('  5.2', 'API Client – axios', 23),
    ('  5.3', 'Các module API', 24),
    ('  5.4', 'Tiện ích Frontend', 25),
    ('  5.5', 'Các trang (Pages)', 25),
    ('  5.6', 'Components dùng chung', 28),
    ('VI.', 'LUỒNG CHỨC NĂNG CHI TIẾT', 29),
    ('  6.1', 'Đăng ký / Đăng nhập', 29),
    ('  6.2', 'Tìm kiếm chuyến tàu', 30),
    ('  6.3', 'Chọn ghế và xem sơ đồ toa', 31),
    ('  6.4', 'Đặt vé (Checkout)', 32),
    ('  6.5', 'Thanh toán QR', 34),
    ('  6.6', 'Tra cứu đơn đặt vé', 36),
    ('  6.7', 'Hủy vé', 37),
    ('  6.8', 'Đổi vé', 38),
    ('  6.9', 'Xem lịch chạy tàu', 39),
    ('  6.10', 'In vé', 40),
    ('VII.', 'BẢO MẬT VÀ XỬ LÝ LỖI', 41),
    ('VIII.', 'CẤU HÌNH MÔI TRƯỜNG', 42),
    ('IX.', 'KẾT LUẬN', 43),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14), leader=2)
    run = p.add_run(f'{num}  {title}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    if num in ('I.','II.','III.','IV.','V.','VI.','VII.','VIII.','IX.'):
        run.bold = True
    run2 = p.add_run(f'\t{page}')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(11)

page_break()

# ═══════════════════════════════════════════════════════════════════
# I. TỔNG QUAN DỰ ÁN
# ═══════════════════════════════════════════════════════════════════
h1('I. TỔNG QUAN DỰ ÁN')
para('KLN Train là hệ thống đặt vé tàu hỏa trực tuyến được xây dựng theo mô hình Full-stack, gồm hai thành phần chính:')
bullet('Frontend: Ứng dụng web SPA (Single Page Application) sử dụng React 19 + Vite, giao diện được thiết kế bằng Tailwind CSS, tương tác người dùng mượt mà với Framer Motion.')
bullet('Backend: REST API xây dựng trên Node.js + Express, kết nối cơ sở dữ liệu SQL Server thông qua Sequelize ORM, xác thực người dùng bằng JWT.')

h2('Công nghệ sử dụng')
add_table(
    ['Thành phần', 'Công nghệ', 'Phiên bản', 'Mục đích'],
    [
        ['Frontend Framework', 'React', '19.2.4', 'Xây dựng giao diện SPA'],
        ['Build Tool', 'Vite', '8.0.4', 'Build & HMR nhanh cho dev'],
        ['CSS Framework', 'Tailwind CSS', '3.4.19', 'Thiết kế UI utility-first'],
        ['Animation', 'Framer Motion', '12.38.0', 'Hiệu ứng chuyển trang, animation'],
        ['HTTP Client', 'Axios', '1.16.1', 'Gọi API từ frontend'],
        ['Routing', 'React Router DOM', '7.14.1', 'Điều hướng trang SPA'],
        ['QR Code', 'qrcode.react', '4.2.0', 'Hiển thị mã QR thanh toán'],
        ['Backend Framework', 'Express', '4.19.2', 'Xây dựng REST API'],
        ['ORM', 'Sequelize', '6.37.8', 'Truy cập SQL Server qua ORM'],
        ['Database Driver', 'mssql + tedious', '11.0.1 / 18.6.1', 'Kết nối SQL Server'],
        ['Authentication', 'jsonwebtoken', '9.0.2', 'Tạo và xác minh JWT'],
        ['Password Hash', 'bcryptjs', '2.4.3', 'Mã hóa mật khẩu'],
        ['Validation', 'express-validator', '7.2.0', 'Kiểm tra đầu vào'],
        ['Database', 'SQL Server (MSSQL)', '2016+', 'Lưu trữ dữ liệu chính'],
        ['Unique ID', 'uuid', '14.0.0', 'Sinh mã giao dịch'],
    ],
    [3.5, 3.5, 2.5, 5]
)

h2('Cấu trúc thư mục gốc')
code_block("""DuAnTauHoa/
├── backend/              ← Node.js REST API
│   ├── app.js            ← Điểm khởi động server
│   ├── package.json
│   ├── .env              ← Biến môi trường (DB, JWT, PORT)
│   └── src/
│       ├── config/       ← Kết nối Sequelize / SQL Server
│       ├── models/       ← 25 Sequelize model + associations
│       ├── repositories/ ← Lớp truy cập dữ liệu (Data Access Layer)
│       ├── services/     ← Lớp nghiệp vụ (Business Logic)
│       ├── controllers/  ← Xử lý HTTP request/response
│       ├── routes/       ← Định tuyến endpoint API
│       ├── middleware/   ← JWT auth, error handler
│       └── utils/        ← Hàm tiện ích (helpers, response)
│
├── frontend/             ← React 19 SPA
│   ├── package.json
│   ├── vite.config.js
│   ├── tailwind.config.js
│   └── src/
│       ├── main.jsx      ← Điểm mount React
│       ├── App.jsx       ← Router cấp cao nhất
│       ├── api/          ← Module gọi API backend
│       ├── pages/        ← 11 trang ứng dụng
│       ├── components/   ← Component dùng chung
│       ├── component/    ← Navbar
│       ├── layout/       ← RootLayout
│       ├── utils/        ← authUtils, dateUtils
│       └── data/         ← Mock data tĩnh (legacy)
│
├── package.json          ← Root (chứa mssql)
└── test_database.js      ← Script kiểm tra kết nối DB""")

page_break()

# ═══════════════════════════════════════════════════════════════════
# II. KIẾN TRÚC TỔNG THỂ
# ═══════════════════════════════════════════════════════════════════
h1('II. KIẾN TRÚC TỔNG THỂ')
para('Hệ thống áp dụng kiến trúc phân tầng (Layered Architecture) ở backend, kết hợp với mô hình SPA ở frontend.')

h2('Sơ đồ tổng thể')
code_block("""┌─────────────────────────────────────────────────────────────────┐
│                     TRÌNH DUYỆT (Browser)                       │
│                    React 19 SPA (Vite)                          │
│   Pages → Components → api/client.js (Axios + JWT interceptor)  │
└──────────────────────────┬──────────────────────────────────────┘
                           │  HTTP/REST  (JSON)
                           │  Authorization: Bearer <JWT>
┌──────────────────────────▼──────────────────────────────────────┐
│                  BACKEND  (Node.js + Express)                    │
│                                                                  │
│  app.js → Routes → Middleware (auth) → Controllers              │
│                                              ↓                   │
│                                         Services                 │
│                                              ↓                   │
│                                        Repositories              │
│                                              ↓                   │
│                                    Sequelize ORM                 │
└──────────────────────────┬──────────────────────────────────────┘
                           │  SQL (TDS Protocol / tedious)
┌──────────────────────────▼──────────────────────────────────────┐
│              SQL SERVER  (KLN_Train Database)                    │
│         25 bảng dữ liệu, quan hệ khóa ngoại, GETUTCDATE()       │
└─────────────────────────────────────────────────────────────────┘""")

h2('Luồng xử lý request')
para('Mỗi request từ trình duyệt đi qua 6 tầng:')
add_table(
    ['Tầng', 'Thành phần', 'Vai trò'],
    [
        ['1', 'React Page / Component', 'Kích hoạt gọi API qua module api/'],
        ['2', 'Axios (api/client.js)', 'Gắn JWT, gửi HTTP request tới backend'],
        ['3', 'Express Routes', 'Khớp URL → gọi middleware → chuyển tới Controller'],
        ['4', 'Middleware (auth.js)', 'Xác minh JWT, gắn req.user hoặc trả 401'],
        ['5', 'Controller', 'Đọc params/body, gọi Service, trả JSON response'],
        ['6', 'Service → Repository → Sequelize → SQL Server', 'Thực thi nghiệp vụ, truy vấn CSDL'],
    ],
    [1.5, 4, 9]
)

h2('Phân tầng Backend (Layered Architecture)')
bullet('Routes: Định nghĩa endpoint, gắn middleware, chuyển tiếp tới Controller tương ứng.')
bullet('Middleware: Xác thực JWT (authenticate, optionalAuth, requireAdmin), xử lý lỗi toàn cục.')
bullet('Controllers: Nhận request, đọc dữ liệu đầu vào (params, query, body), gọi Service, trả response chuẩn hóa.')
bullet('Services: Chứa toàn bộ logic nghiệp vụ. Không biết về HTTP. Gọi Repository và thực thi transaction Sequelize.')
bullet('Repositories: Lớp duy nhất tương tác trực tiếp với model Sequelize. Cung cấp các phương thức query cụ thể.')
bullet('Models: Ánh xạ bảng SQL Server sang class JavaScript. Khai báo associations (quan hệ) trong index.js.')

page_break()

# ═══════════════════════════════════════════════════════════════════
# III. CƠ SỞ DỮ LIỆU
# ═══════════════════════════════════════════════════════════════════
h1('III. CƠ SỞ DỮ LIỆU – SQL SERVER KLN_TRAIN')

h2('3.1. Danh sách bảng và mục đích')
add_table(
    ['STT', 'Tên bảng', 'Nhóm', 'Mục đích'],
    [
        ['1', 'TaiKhoan', 'Người dùng', 'Tài khoản đăng nhập (khách hàng / quản trị)'],
        ['2', 'HanhKhach', 'Người dùng', 'Hành khách đi tàu (liên kết tài khoản)'],
        ['3', 'Tau', 'Hạ tầng tàu', 'Thông tin đoàn tàu (số hiệu, số toa)'],
        ['4', 'GaTau', 'Hạ tầng tàu', 'Danh sách ga tàu trên tuyến'],
        ['5', 'LoaiToa', 'Hạ tầng tàu', 'Loại toa (toa cứng, mềm, nằm…)'],
        ['6', 'LoaiGhe', 'Hạ tầng tàu', 'Loại ghế trong từng loại toa + hệ số giá'],
        ['7', 'CauHinhToa', 'Hạ tầng tàu', 'Cấu hình toa của mỗi đoàn tàu (toa 1,2,3…)'],
        ['8', 'CauHinhGhe', 'Hạ tầng tàu', 'Sơ đồ ghế vật lý trong từng loại toa'],
        ['9', 'LichChay', 'Lịch trình', 'Lịch chạy tàu (tuyến, giờ, ngày trong tuần)'],
        ['10', 'LichTrinhChuyen', 'Lịch trình', 'Các ga dừng dọc đường + khoảng cách km'],
        ['11', 'ChuyenTau', 'Lịch trình', 'Chuyến tàu cụ thể theo ngày chạy'],
        ['12', 'BieuGia', 'Giá vé', 'Biểu giá theo đợt (hè, Tết, thường)'],
        ['13', 'ChinhSachGia', 'Giá vé', 'Chính sách giảm giá theo loại hành khách'],
        ['14', 'KhuyenMai', 'Giá vé', 'Mã khuyến mãi (% hoặc số tiền cố định)'],
        ['15', 'ChinhSachHuy', 'Chính sách', 'Phí hủy vé theo số giờ trước giờ chạy'],
        ['16', 'DonDatVe', 'Đặt vé', 'Đơn đặt vé (gom nhiều vé, thông tin liên lạc)'],
        ['17', 'Ve', 'Đặt vé', 'Vé cụ thể cho từng hành khách, từng chuyến'],
        ['18', 'TamGiuGhe', 'Đặt vé', 'Giữ ghế tạm thời 15 phút trong quá trình checkout'],
        ['19', 'ThanhToan', 'Thanh toán', 'Giao dịch thanh toán (QR VietQR)'],
        ['20', 'HoaDon', 'Thanh toán', 'Hóa đơn xuất sau khi thanh toán thành công'],
        ['21', 'HoanTien', 'Hủy/Đổi', 'Thông tin hoàn tiền khi hủy vé'],
        ['22', 'DoiVe', 'Hủy/Đổi', 'Bản ghi đổi vé (vé cũ → vé mới + phí)'],
        ['23', 'DonKhuHoi', 'Đặt vé', 'Liên kết đơn chiều đi và chiều về (vé khứ hồi)'],
        ['24', 'PhanHoi', 'Hệ thống', 'Đánh giá, phản hồi của khách hàng sau chuyến'],
        ['25', 'ThongBao', 'Hệ thống', 'Thông báo hệ thống gửi đến tài khoản'],
        ['26', 'AuditLog', 'Hệ thống', 'Nhật ký thay đổi dữ liệu (ai, khi nào, thay đổi gì)'],
    ],
    [0.8, 3.5, 3, 7.7]
)

h2('3.2. Mô tả chi tiết từng bảng')

h3('Bảng TaiKhoan – Tài khoản người dùng')
add_table(
    ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['id_tai_khoan', 'INT', 'PK, AUTO', 'Khóa chính'],
        ['email', 'NVARCHAR(100)', 'UNIQUE, NOT NULL', 'Email đăng nhập'],
        ['mat_khau', 'VARCHAR(255)', 'NOT NULL', 'Mật khẩu đã hash bcrypt (10 rounds)'],
        ['ho_ten', 'NVARCHAR(100)', 'NOT NULL', 'Họ và tên'],
        ['so_dien_thoai', 'VARCHAR(15)', '', 'Số điện thoại'],
        ['ngay_sinh', 'DATE', '', 'Ngày sinh'],
        ['gioi_tinh', 'VARCHAR(10)', '', 'nam / nu / khac'],
        ['vai_tro', 'VARCHAR(20)', 'DEFAULT khach_hang', 'khach_hang | quan_tri'],
        ['trang_thai', 'VARCHAR(20)', 'DEFAULT hoat_dong', 'hoat_dong | khoa'],
        ['ngay_tao', 'DATETIME', 'DEFAULT NOW', 'Thời điểm tạo tài khoản'],
    ],
    [3.5, 3.5, 3.5, 4.5]
)

h3('Bảng Ve – Vé tàu')
para('Bảng trung tâm của hệ thống đặt vé. Mỗi vé tương ứng 1 hành khách, 1 chuyến tàu, 1 ghế.')
add_table(
    ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['id_ve', 'INT', 'PK, AUTO', 'Khóa chính'],
        ['id_don_dat_ve', 'INT', 'FK → DonDatVe', 'Thuộc đơn đặt vé nào'],
        ['id_hanh_khach', 'INT', 'FK → HanhKhach', 'Hành khách đi'],
        ['id_chuyen', 'INT', 'FK → ChuyenTau', 'Chuyến tàu cụ thể'],
        ['so_toa_thu_tu', 'INT', 'NOT NULL', 'Số thứ tự toa (1,2,3…)'],
        ['so_ghe_trong_toa', 'INT', 'NOT NULL', 'Số ghế trong toa'],
        ['id_ga_len', 'INT', 'FK → GaTau', 'Ga lên (ga xuất phát của hành khách)'],
        ['id_ga_xuong', 'INT', 'FK → GaTau', 'Ga xuống (ga đến của hành khách)'],
        ['loai_hanh_khach', 'VARCHAR(20)', '', 'nguoi_lon | tre_em'],
        ['gia_ve', 'DECIMAL(12,2)', 'NOT NULL', 'Giá vé sau tính toán'],
        ['qr_ve', 'TEXT', '', 'Chuỗi QR code khi in vé'],
        ['trang_thai', 'VARCHAR(20)', '', 'cho_xac_nhan → da_xac_nhan → da_huy | da_doi'],
        ['ngay_xuat_ve', 'DATETIME', '', 'Thời điểm xuất vé'],
        ['id_cs_huy', 'INT', 'FK → ChinhSachHuy', 'Chính sách hủy đã áp dụng (nếu hủy)'],
    ],
    [3.5, 3.5, 3.5, 4.5]
)

h3('Bảng TamGiuGhe – Giữ ghế tạm thời')
para('Cơ chế ngăn chặn đặt trùng ghế trong quá trình thanh toán. Mỗi bản ghi có hiệu lực 15 phút.')
add_table(
    ['Cột', 'Kiểu dữ liệu', 'Mô tả'],
    [
        ['id_giu', 'INT PK', 'Khóa chính'],
        ['id_chuyen', 'INT FK', 'Chuyến tàu đang giữ ghế'],
        ['so_toa_thu_tu', 'INT', 'Số thứ tự toa'],
        ['so_ghe_trong_toa', 'INT', 'Số ghế đang giữ'],
        ['session_id', 'VARCHAR(20)', 'ID phiên checkout (ngẫu nhiên 8 byte hex)'],
        ['id_don_dat_ve', 'INT FK', 'Liên kết đơn sau khi tạo (NULL khi chỉ giữ)'],
        ['trang_thai', 'VARCHAR(20)', 'dang_giu → da_dat | da_giai_phong'],
        ['thoi_gian_giu', 'DATETIME', 'Thời điểm bắt đầu giữ'],
        ['thoi_gian_het_han', 'DATETIME', 'Thời điểm hết hiệu lực (= thoi_gian_giu + 15 phút)'],
    ],
    [3.5, 3, 8.5]
)

h3('Bảng BieuGia – Biểu giá')
para('Quy định giá vé theo từng đợt (hè, Tết, thường). Công thức tính: Giá = km × don_gia_km_goc × he_so_tang × he_so_ghe, làm tròn lên 1.000đ.')
add_table(
    ['Cột', 'Mô tả'],
    [
        ['id_bieu_gia', 'Khóa chính'],
        ['ten_dip', 'Tên đợt (Hè 2026, Tết 2026, Thường…)'],
        ['don_gia_km_goc', 'Đơn giá gốc mỗi km (VD: 264 đ/km)'],
        ['he_so_tang', 'Hệ số tăng theo mùa (1.0, 1.2, 1.5…)'],
        ['ngay_bat_dau / ngay_ket_thuc', 'Phạm vi áp dụng'],
        ['trang_thai', 'dang_ap_dung | het_han'],
    ],
    [4, 11]
)

h3('Bảng ChinhSachHuy – Chính sách hủy vé')
para('Quy định phí hủy dựa trên số giờ còn lại trước giờ tàu chạy. Lưu trong DB để quản trị viên có thể điều chỉnh.')
add_table(
    ['Giờ còn lại', 'Phí hủy (%)', 'Tiền hoàn (%)'],
    [
        ['≥ 72 giờ (3 ngày)', '10%', '90%'],
        ['24 – 72 giờ (1–3 ngày)', '25%', '75%'],
        ['4 – 24 giờ', '50%', '50%'],
        ['< 4 giờ hoặc sau giờ chạy', '100%', '0%'],
    ],
    [5, 4, 4]
)

h2('3.3. Quan hệ giữa các bảng (Associations)')
code_block("""TaiKhoan ──< HanhKhach      (1 tài khoản có nhiều hành khách)
TaiKhoan ──< DonDatVe       (1 tài khoản có nhiều đơn đặt vé)
TaiKhoan ──< ThongBao       (thông báo gửi cho tài khoản)

Tau ──< CauHinhToa >── LoaiToa   (tàu có nhiều toa, mỗi toa là 1 loại)
LoaiToa ──< LoaiGhe             (loại toa có nhiều loại ghế)
LoaiToa ──< CauHinhGhe          (sơ đồ vật lý ghế theo loại toa)

Tau ──< LichChay                (tàu có nhiều lịch chạy)
LichChay >── GaTau (ga đi)      (lịch chạy có ga xuất phát)
LichChay >── GaTau (ga đến)     (lịch chạy có ga đích)
LichChay ──< LichTrinhChuyen >── GaTau  (các ga dừng dọc đường)
LichChay ──< ChuyenTau          (lịch chạy có nhiều chuyến theo ngày)

DonDatVe ──< Ve                 (1 đơn có nhiều vé)
DonDatVe ──< ThanhToan          (1 đơn có thể có nhiều giao dịch)
DonDatVe ──< TamGiuGhe          (giữ ghế trong quá trình đặt)
DonDatVe >── KhuyenMai          (áp mã giảm giá nếu có)
ThanhToan ──1 HoaDon            (1 giao dịch → 1 hóa đơn)

Ve >── HanhKhach                (vé thuộc về hành khách)
Ve >── ChuyenTau                (vé của chuyến tàu nào)
Ve >── GaTau (ga lên)
Ve >── GaTau (ga xuống)

Ve ──1 HoanTien                 (nếu hủy → tạo bản ghi hoàn tiền)
DoiVe >── Ve (vé cũ)
DoiVe >── Ve (vé mới)

DonKhuHoi >── DonDatVe (đơn đi)
DonKhuHoi >── DonDatVe (đơn về)  (vé khứ hồi)""")

page_break()

# ═══════════════════════════════════════════════════════════════════
# IV. BACKEND
# ═══════════════════════════════════════════════════════════════════
h1('IV. BACKEND – NODE.JS / EXPRESS')

h2('4.1. Điểm khởi động: backend/app.js')
para('File app.js là điểm khởi động toàn bộ backend. Nó thực hiện các bước sau theo thứ tự:')
bullet('Gọi dotenv.config() để nạp biến môi trường từ file .env.')
bullet('Khởi tạo Express app, cấu hình CORS chỉ cho phép origin từ FRONTEND_URL (mặc định http://localhost:5173).')
bullet('Thêm middleware express.json({ limit: "10mb" }) và express.urlencoded() để parse request body.')
bullet('Mount toàn bộ routes tại /api thông qua require("./src/routes").')
bullet('Gắn notFoundHandler (trả 404) và errorHandler (trả lỗi chuẩn hóa) vào cuối pipeline.')
bullet('Gọi connectDB() để xác thực kết nối SQL Server, sau đó app.listen() trên PORT (mặc định 5000).')
bullet('Require("./src/models") để đảm bảo tất cả Sequelize associations được khai báo khi server khởi động.')

code_block("""// Luồng khởi động:
dotenv.config()  →  express()  →  cors()  →  json()  →  routes('/api')
→  notFoundHandler  →  errorHandler  →  connectDB()  →  app.listen(5000)""")

h2('4.2. Cấu hình CSDL: src/config/database.js')
para('File này khởi tạo instance Sequelize kết nối tới SQL Server Named Instance (DESKTOP-5DF60PC\\SQLEXPRESS):')
bullet('Đọc DB_SERVER từ .env, tách thành host (tên máy) và instanceName (SQLEXPRESS) bằng dấu \\\\.')
bullet('Sử dụng dialect: "mssql" với driver tedious qua dialectOptions.')
bullet('Tắt mã hóa (encrypt: false) và bật trustServerCertificate: true cho môi trường nội bộ.')
bullet('Cấu hình connection pool: tối đa 10 kết nối, idle timeout 10 giây.')
bullet('Chỉ bật logging (console.log SQL) khi NODE_ENV === "development".')
bullet('Export hàm connectDB() gọi sequelize.authenticate() để kiểm tra kết nối khi khởi động.')

h2('4.3. Models – Sequelize ORM (src/models/)')
para('Có tổng cộng 25 model, mỗi model ánh xạ 1:1 với bảng trong SQL Server. Tất cả associations được khai báo tập trung trong index.js.')

h3('Cấu trúc một Model điển hình')
code_block("""const { DataTypes } = require('sequelize')
const { sequelize } = require('../config/database')

const TaiKhoan = sequelize.define('TaiKhoan', {
  id_tai_khoan:  { type: DataTypes.INTEGER, primaryKey: true, autoIncrement: true },
  email:         { type: DataTypes.STRING(100), unique: true, allowNull: false },
  mat_khau:      { type: DataTypes.STRING(255), allowNull: false },
  vai_tro:       { type: DataTypes.STRING(20), defaultValue: 'khach_hang' },
  trang_thai:    { type: DataTypes.STRING(20), defaultValue: 'hoat_dong' },
  ngay_tao:      { type: DataTypes.DATE, defaultValue: DataTypes.NOW },
}, { tableName: 'TaiKhoan', timestamps: false })

module.exports = TaiKhoan""")

h3('Khai báo Associations trong models/index.js')
para('File index.js require tất cả 25 model rồi khai báo các mối quan hệ. Điều này quan trọng vì Sequelize dùng associations để sinh câu JOIN khi dùng include trong query. Ví dụ:')
code_block("""// Quan hệ 1:N – TaiKhoan có nhiều HanhKhach
TaiKhoan.hasMany(HanhKhach, { foreignKey: 'id_tai_khoan' })
HanhKhach.belongsTo(TaiKhoan, { foreignKey: 'id_tai_khoan' })

// Quan hệ đa chiều – Ve thuộc GaTau theo 2 vai trò (alias)
Ve.belongsTo(GaTau, { foreignKey: 'id_ga_len',   as: 'GaLen'   })
Ve.belongsTo(GaTau, { foreignKey: 'id_ga_xuong', as: 'GaXuong' })

// Quan hệ LichChay ↔ GaTau (2 khóa ngoại, cần alias)
LichChay.belongsTo(GaTau, { foreignKey: 'id_ga_di',  as: 'GaDi'  })
LichChay.belongsTo(GaTau, { foreignKey: 'id_ga_den', as: 'GaDen' })""")

h2('4.4. Repositories – Lớp truy cập dữ liệu (src/repositories/)')
para('Repository là lớp duy nhất được phép gọi trực tiếp Sequelize. Service không dùng Model trực tiếp mà gọi qua Repository (ngoại trừ một số Service dùng thêm sequelize.query() cho raw SQL).')

h3('BaseRepository.js – CRUD dùng chung')
add_table(
    ['Phương thức', 'Mô tả'],
    [
        ['findAll(options)', 'Model.findAll() với options tùy ý (where, include, order…)'],
        ['findById(id, options)', 'Model.findByPk(id) – tìm theo khóa chính'],
        ['findOne(where, options)', 'Model.findOne({ where }) – tìm 1 bản ghi'],
        ['create(data, options)', 'Model.create(data) – tạo mới'],
        ['update(id, data, pkField)', 'Model.update(data, { where: {pkField: id} }) – cập nhật theo PK'],
        ['updateWhere(where, data)', 'Model.update(data, { where }) – cập nhật hàng loạt'],
        ['delete(id, pkField)', 'Model.destroy({ where: {pkField: id} })'],
        ['count(where)', 'Model.count({ where }) – đếm bản ghi'],
        ['findAndCountAll(options)', 'Model.findAndCountAll() – phân trang'],
    ],
    [5, 10]
)

h3('TaiKhoanRepository.js')
bullet('extends BaseRepository với model TaiKhoan.')
bullet('findByEmail(email): tìm tài khoản theo email, trim và case-insensitive.')
bullet('emailExists(email): kiểm tra email đã tồn tại chưa (trả về boolean).')

h3('TrainRepository.js')
para('Repository phức tạp nhất, xử lý các query liên quan đến tàu, lịch chạy, sơ đồ ghế.')
add_table(
    ['Hàm', 'Query thực hiện', 'Mô tả'],
    [
        ['searchChuyen(idGaDi, idGaDen, ngayChay)',
         'ChuyenTau.findAll với include: LichChay (where id_ga_di, id_ga_den), Tau, GaTau (GaDi/GaDen). ORDER BY gio_khoi_hanh ASC.',
         'Tìm chuyến tàu theo tuyến và ngày'],
        ['getCoachesByChuyen(idChuyen)',
         'Tìm ChuyenTau → lấy id_tau → CauHinhToa.findAll include LoaiToa, ORDER BY so_toa_thu_tu.',
         'Danh sách toa của chuyến'],
        ['getSeatMap(idChuyen, soToaThuTu)',
         '3 query riêng: (1) CauHinhGhe.findAll include LoaiGhe; (2) Ve.findAll (not da_huy/da_doi); (3) TamGiuGhe.findAll (dang_giu, het_han > now). Merge thành trạng thái empty/sold/held.',
         'Sơ đồ ghế với trạng thái thực tế'],
        ['getLichTrinh(idLichChay)',
         'LichTrinhChuyen.findAll include GaTau, ORDER BY thu_tu_dung.',
         'Các ga dừng dọc đường'],
        ['getAllGa()',
         'GaTau.findAll where trang_thai=hoat_dong, ORDER BY do_uu_tien DESC.',
         'Danh sách ga đang hoạt động'],
    ],
    [4, 6.5, 4.5]
)

h3('BookingRepository.js')
add_table(
    ['Hàm', 'Mô tả'],
    [
        ['findByMaDon(maDon)', 'Tìm đơn theo mã đơn, include đầy đủ: Ve → HanhKhach, ChuyenTau → LichChay → Tau/GaTau, ThanhToan, KhuyenMai.'],
        ['findByMaDatCho(maDatCho)', 'Tương tự nhưng tìm theo mã đặt chỗ (toUpperCase + trim).'],
        ['findByTaiKhoan(idTaiKhoan)', 'Tất cả đơn của người dùng, ORDER BY thoi_gian_dat DESC.'],
        ['holdSeats(idChuyen, soToa, soGheList, sessionId, idDon)', 'Xóa TamGiuGhe hết hạn trước, rồi TamGiuGhe.bulkCreate() các ghế với thoi_gian_het_han = now+15min.'],
        ['releaseHoldBySession(sessionId)', 'TamGiuGhe.update trạng thái → da_giai_phong cho session.'],
        ['checkSeatsAvailable(idChuyen, soToa, soGheList, excludeSessionId)', 'Ve.count (not huy/doi) + TamGiuGhe.count (dang_giu, chưa hết hạn, không phải session bị exclude). Trả true nếu cả 2 = 0.'],
    ],
    [5, 10]
)

h2('4.5. Services – Lớp nghiệp vụ (src/services/)')

h3('AuthService.js')
add_table(
    ['Hàm', 'Logic thực hiện'],
    [
        ['register({email, matKhau, hoTen, soDienThoai})',
         '1. emailExists() → ném lỗi nếu đã tồn tại\n2. bcrypt.hash(matKhau, 10) → tạo hash\n3. TaiKhoan.create() với vai_tro=khach_hang\n4. signToken() → trả { token, user }'],
        ['login({email, matKhau})',
         '1. findByEmail() → ném 401 nếu không tìm thấy\n2. Kiểm tra trang_thai !== "khoa"\n3. bcrypt.compare() → ném 401 nếu sai mật khẩu\n4. signToken() → trả { token, user }'],
        ['signToken(taiKhoan)',
         'jwt.sign({ id, email, role }, JWT_SECRET, { expiresIn: "24h" })'],
        ['getProfile(idTaiKhoan)',
         'findById() → sanitize() loại bỏ trường mat_khau → camelCase'],
        ['changePassword(id, {matKhauCu, matKhauMoi})',
         '1. findById()\n2. bcrypt.compare() với mật khẩu cũ\n3. bcrypt.hash(matKhauMoi, 10)\n4. update() mật khẩu mới'],
    ],
    [4, 11]
)

h3('TrainService.js')
para('Service quan trọng nhất, xử lý tìm kiếm tàu và tính giá vé.')
add_table(
    ['Hàm', 'Logic'],
    [
        ['tinhGiaVe(idLichChay, ngayChay, idGaLen, idGaXuong, idLoaiGhe, hesoGhe)',
         '1. Raw SQL: SELECT khoang_cach_km FROM LichTrinhChuyen WHERE id_lich_chay=? AND id_ga=? (x2 lần cho ga lên và ga xuống)\n2. km = |km_xuong - km_len|\n3. BieuGia.findAll (ngay_bat_dau ≤ ngayChay ≤ ngay_ket_thuc, dang_ap_dung) → lấy don_gia_km_goc + he_so_tang\n4. Công thức: Math.ceil(km × donGia × hesoTang × hesoGhe / 1000) × 1000'],
        ['searchTrains(tenGaDi, tenGaDen, ngayChay)',
         '1. GaTau.findOne x2 (tìm ga theo tên)\n2. TrainRepo.searchChuyen() → danh sách chuyến\n3. Với mỗi chuyến: getCoachesByChuyen() + tinhGiaVe(hesoGhe=1.0) → priceFrom\n4. Trả array thông tin chuyến kèm danh sách toa'],
        ['getSeatMap(idChuyen, soToa, idGaLen, idGaXuong)',
         '1. TrainRepo.getSeatMap() → layout ghế với trạng thái\n2. Nếu có ga: tinhGiaVe() cho từng loại ghế, dùng priceCache tránh query lặp\n3. Trả seats với trường gia bổ sung'],
        ['getTrainRouteDetail(idLichChay, idGaLen, idGaXuong, ngayChay)',
         '1. LichChay.findByPk include Tau → CauHinhToa → LoaiToa → LoaiGhe\n2. getLichTrinh() → danh sách ga dừng\n3. Gom loaiGheMap từ tất cả toa\n4. tinhGiaVe() cho từng loại ghế → priceRows'],
    ],
    [4.5, 10.5]
)

h3('BookingService.js')
para('Service phức tạp nhất, xử lý toàn bộ quy trình đặt vé trong một Sequelize transaction.')
add_table(
    ['Bước', 'Thao tác trong createBooking()'],
    [
        ['1. Validate khuyến mãi', 'KhuyenMai.findOne → kiểm tra ngày hết hạn + số lượt dùng'],
        ['2. Tính tiền', 'tongTienVe = Σ(seatPrice); phiDichVu = 20k (1 chiều) / 40k (khứ hồi); áp tienGiam từ khuyến mãi'],
        ['3. Kiểm tra ghế', 'checkSeatsAvailable() – nhóm theo toa; cho phép exclude session đang giữ'],
        ['4. Sinh mã', 'genBookingCode() (6 ký tự) vòng lặp cho đến khi unique; genOrderCode() (KLN+6 số)'],
        ['5. Tạo DonDatVe', 'DonDatVe.create() với thoi_gian_het_han = now+15min'],
        ['6. Tạo HanhKhach + Ve', 'Với mỗi hành khách: HanhKhach.findOne (theo tên+ngaySinh) hoặc create; Ve.create()'],
        ['7. Giữ ghế', 'Nếu có sessionId: TamGiuGhe.update() link vào đơn; nếu không: holdSeats() tạo mới'],
        ['8. Cập nhật KM', 'KhuyenMai.increment("da_dung", { by: 1 })'],
        ['Trả về', '{ don, maDon, maDatCho, tongThanhToan, tienGiam, veList }'],
    ],
    [4, 11]
)

h3('PaymentService.js')
add_table(
    ['Hàm', 'Logic'],
    [
        ['createPayment(idDonDatVe, phuongThuc)',
         '1. Kiểm tra đơn tồn tại, trạng thái = cho_thanh_toan, chưa hết hạn 15 phút\n2. genTransactionCode() → UUID 20 ký tự\n3. Tạo URL VietQR: https://img.vietqr.io/image/BIDV-...?amount=X&addInfo=maDon\n4. ThanhToan.create() → trả { idThanhToan, maGiaoDich, qrUrl, soTien }'],
        ['confirmPayment(idThanhToan)',
         'TRANSACTION:\n1. ThanhToan.update → thanh_cong\n2. DonDatVe.update → da_thanh_toan\n3. Ve.update → da_xac_nhan (all cho_xac_nhan)\n4. TamGiuGhe.update → da_dat\n5. genInvoiceNumber() → HoaDon.create()'],
    ],
    [4.5, 10.5]
)

h3('CancelService.js')
add_table(
    ['Hàm', 'Logic'],
    [
        ['tinhPhiHuy(idVe)',
         '1. Ve.findByPk include ChuyenTau → LichChay\n2. departAt = new Date(ngayChay + gioKhoiHanh)\n3. gioConLai = (departAt - now) / 3.6M ms\n4. ChinhSachHuy.findOne (gio_truoc_gio_chay ≥ floor(gioConLai), ORDER ASC) → phiHuyPct\n5. phiHuy = giaVe × phiHuyPct / 100; tienHoan = giaVe - phiHuy'],
        ['cancelTickets(maDatCho, idVeList, lyDo)',
         'TRANSACTION:\n1. Tìm đơn + xác nhận đã thanh toán\n2. Với mỗi vé: tinhPhiHuy() → ve.update(da_huy) → HoanTien.create()\n3. Kiểm tra tất cả vé đã hủy → DonDatVe.update(da_huy)\n4. Trả { tongTienHoan, soVeHuy }'],
    ],
    [4.5, 10.5]
)

h3('ExchangeService.js')
add_table(
    ['Hàm', 'Logic'],
    [
        ['checkExchangeable(idVe)',
         '1. Ve.findByPk include ChuyenTau → LichChay\n2. Kiểm tra trang_thai = da_xac_nhan\n3. gioConLai ≥ 24h → mới cho phép đổi\n4. calcExchangeFee(giaVe) = max(giaVe×5%, 20000)\n5. Trả { idVe, giaVe, phiDoi, gioConLai }'],
        ['exchangeTicket(idVeCu, newTicketData)',
         'TRANSACTION:\n1. checkSeatsAvailable() cho ghế mới\n2. Ve.create() → vé mới với cùng hành khách\n3. phiDoi = calcExchangeFee(); chenhLech = max(0, giaVeMoi-giaVeCu)\n4. DoiVe.create({ id_ve_cu, id_ve_moi, phi_doi, chenh_lech, tong_phai_tra })\n5. veCu.update(da_doi)'],
    ],
    [4.5, 10.5]
)

h2('4.6. Controllers – Xử lý HTTP Request (src/controllers/)')
para('Mỗi Controller nhận request từ Express, đọc params/query/body, gọi Service tương ứng, trả response chuẩn hóa qua utils/response.js.')

h3('AuthController.js')
add_table(
    ['Endpoint', 'Method + Path', 'Xử lý'],
    [
        ['register', 'POST /api/auth/register', 'validationResult() → AuthService.register() → created(res, {token,user})'],
        ['login', 'POST /api/auth/login', 'validationResult() → AuthService.login() → ok(res, {token,user})'],
        ['getProfile', 'GET /api/auth/profile [Auth]', 'req.user.id → AuthService.getProfile() → ok(res, user)'],
        ['updateProfile', 'PUT /api/auth/profile [Auth]', 'req.body → AuthService.updateProfile() → ok(res, user)'],
        ['changePassword', 'PUT /api/auth/change-password [Auth]', 'req.body → AuthService.changePassword() → ok(res)'],
    ],
    [3.5, 4.5, 7]
)

h3('TrainController.js')
add_table(
    ['Endpoint', 'Method + Path', 'Xử lý'],
    [
        ['searchTrains', 'GET /api/trains/search', 'req.query: gaDi, gaDen, ngayChay → TrainService.searchTrains()'],
        ['getSeatMap', 'GET /api/trains/:idChuyen/seats/:soToa', 'params + query (idGaLen, idGaXuong) → TrainService.getSeatMap()'],
        ['getAllStations', 'GET /api/trains/stations', 'TrainService.getAllStations()'],
        ['getSchedule', 'GET /api/trains/schedule', 'TrainService.getSchedule()'],
        ['getTrainDetail', 'GET /api/trains/detail/:idLichChay', 'params + query → TrainService.getTrainRouteDetail()'],
    ],
    [3.5, 5, 6.5]
)

h3('BookingController.js')
add_table(
    ['Endpoint', 'Method + Path', 'Auth', 'Xử lý'],
    [
        ['holdSeats', 'POST /api/bookings/hold-seats', '—', 'req.body.trips → BookingService.holdSeatsForCheckout() → {sessionId, hetHan}'],
        ['createBooking', 'POST /api/bookings', 'OptionalAuth', 'req.body → BookingService.createBooking() → created(res, {maDon,maDatCho,idDon,tongThanhToan})'],
        ['lookupBooking', 'POST /api/bookings/lookup', '—', '{maDatCho,email,phone} → BookingService.lookupBooking() → formatDon()'],
        ['getBookingHistory', 'GET /api/bookings/history', 'Auth', 'req.user.id → BookingService.getBookingHistory()'],
        ['getBookingByCode', 'GET /api/bookings/:maDatCho', '—', 'BookingRepo.findByMaDatCho() → formatDon()'],
    ],
    [3, 4.5, 1.8, 5.7]
)

h2('4.7. Routes – Định tuyến API (src/routes/)')
para('File index.js gộp tất cả router con và mount vào app, đồng thời định nghĩa GET /api/health.')
code_block("""router.use('/auth',     authRoutes)
router.use('/trains',   trainRoutes)
router.use('/bookings', bookingRoutes)
router.use('/payments', paymentRoutes)
router.use('/cancel',   cancelRoutes)
router.use('/exchange', exchangeRoutes)
router.get('/health', (req, res) => res.json({ status: 'OK', time: new Date() }))""")

h2('4.8. Middleware (src/middleware/)')
h3('auth.js – Xác thực JWT')
add_table(
    ['Middleware', 'Hành vi'],
    [
        ['authenticate', 'Yêu cầu header Authorization: Bearer <token>. Verify JWT → gắn req.user = {id, email, role}. Trả 401 nếu thiếu hoặc token không hợp lệ/hết hạn.'],
        ['optionalAuth', 'Tương tự authenticate nhưng không ném lỗi nếu thiếu token. req.user = null nếu không có hoặc token lỗi.'],
        ['requireAdmin', 'Kiểm tra req.user.role === "quan_tri". Trả 403 nếu không phải admin. Luôn dùng sau authenticate.'],
    ],
    [3, 12]
)

h3('errorHandler.js – Xử lý lỗi toàn cục')
bullet('notFoundHandler: Bắt mọi URL không khớp với route nào → trả 404.')
bullet('errorHandler: Bắt lỗi ném ra từ next(err) hoặc Service. Xử lý Sequelize validation error (400), Sequelize unique constraint (409), lỗi tùy chỉnh với err.status, lỗi không xác định (500).')

h2('4.9. Utils – Tiện ích (src/utils/)')
h3('helpers.js')
add_table(
    ['Hàm', 'Mô tả'],
    [
        ['genBookingCode()', '6 ký tự ngẫu nhiên từ "ABCDEFGHJKLMNPQRSTUVWXYZ23456789" (loại bỏ I, O, L, 0, 1 dễ nhầm)'],
        ['genOrderCode()', '"KLN" + 6 chữ số ngẫu nhiên (100000–999999)'],
        ['genTransactionCode()', 'uuidv4() bỏ dấu gạch, lấy 20 ký tự đầu, UPPERCASE'],
        ['genInvoiceNumber()', '"HD" + YYYYMMDD + 4 chữ số ngẫu nhiên (0000–9999)'],
        ['formatVND(amount)', 'Intl.NumberFormat("vi-VN").format() + " đ"'],
        ['calcExchangeFee(originalPrice)', 'max(originalPrice × 5%, 20000) – phí đổi vé tối thiểu 20.000đ'],
    ],
    [4.5, 10.5]
)

h3('response.js')
code_block("""ok(res, data, message, statusCode=200)     → { success:true, message, data }
created(res, data, message)                 → status 201
error(res, message, statusCode=500, errors) → { success:false, message, errors }
notFound(res, message)                      → status 404
badRequest(res, message, errors)            → status 400
unauthorized(res, message)                  → status 401
forbidden(res, message)                     → status 403""")

page_break()

# ═══════════════════════════════════════════════════════════════════
# V. FRONTEND
# ═══════════════════════════════════════════════════════════════════
h1('V. FRONTEND – REACT 19 / VITE')

h2('5.1. Cấu trúc và điểm khởi động')
h3('main.jsx – Điểm mount React')
code_block("""import { StrictMode } from 'react'
import { createRoot } from 'react-dom/client'
import App from './App.jsx'
import './index.css'
createRoot(document.getElementById('root')).render(<StrictMode><App /></StrictMode>)""")

h3('App.jsx – Router cấp cao nhất')
para('App.jsx định nghĩa toàn bộ routing của SPA, chia thành 2 nhóm:')
add_table(
    ['Path', 'Component', 'Layout'],
    [
        ['/', 'Home', 'MainLayout (có Navbar)'],
        ['/tim-ve', 'TicketSearch', 'MainLayout'],
        ['/checkout', 'Checkout', 'MainLayout'],
        ['/thanh-toan', 'PaymentMethod', 'MainLayout'],
        ['/thanh-toan/qr', 'QRPayment', 'MainLayout'],
        ['/thanh-toan/thanh-cong', 'PaymentSuccess', 'MainLayout'],
        ['/thong-tin-dat-cho', 'BookingLookup', 'MainLayout'],
        ['/chuyen-tau-gia-ve', 'TrainSchedule', 'MainLayout'],
        ['/tra-ve', 'CancelTicket', 'MainLayout'],
        ['/doi-ve', 'ExchangeTicket', 'MainLayout'],
        ['/in-ve', 'PrintTicket', 'Không có Navbar'],
    ],
    [4, 4, 7]
)

para('MainLayout là component wrapper render <Navbar /> + <Outlet /> (nội dung trang con). Sử dụng React Router Outlet pattern.')

h2('5.2. API Client – src/api/client.js')
para('Được xây dựng trên Axios với hai interceptors quan trọng:')
h3('Request Interceptor')
bullet('Đọc localStorage["KLN_AUTH"] → parse JSON → lấy token.')
bullet('Gắn Authorization: Bearer <token> vào header nếu có.')
bullet('Log request (method, URL, body) trong môi trường development.')

h3('Response Interceptor')
bullet('Trả về response.data trực tiếp (bỏ qua wrapper axios) → caller nhận JSON body ngay.')
bullet('Xử lý lỗi mạng (ERR_NETWORK): ném Error với message hướng dẫn kiểm tra backend.')
bullet('Xử lý timeout (ECONNABORTED): thông báo quá thời gian 20 giây.')
bullet('Xử lý HTTP error (4xx, 5xx): normalize thành Error với status và errors từ response.data.')
bullet('Log chi tiết lỗi theo loại trong dev mode để debug.')

h2('5.3. Các module API (src/api/)')
add_table(
    ['File', 'Hàm xuất', 'Endpoint gọi'],
    [
        ['auth.js', 'register(data)\nlogin(data)\ngetProfile()\nupdateProfile(data)\nchangePassword(data)',
         'POST /auth/register\nPOST /auth/login\nGET /auth/profile\nPUT /auth/profile\nPUT /auth/change-password'],
        ['trains.js', 'searchTrains(gaDi, gaDen, ngayChay)\ngetStations()\ngetSchedule()\ngetSeatMap(idChuyen, soToa, idGaLen, idGaXuong)\ngetTrainDetail(...)',
         'GET /trains/search?...\nGET /trains/stations\nGET /trains/schedule\nGET /trains/:id/seats/:soToa?...\nGET /trains/detail/:id?...'],
        ['bookings.js', 'holdSeats(data)\ncreateBooking(data)\nlookupBooking(maDatCho, email, phone)\ngetBookingHistory()\ngetBookingByCode(maDatCho)',
         'POST /bookings/hold-seats\nPOST /bookings\nPOST /bookings/lookup\nGET /bookings/history\nGET /bookings/:code'],
        ['payments.js', 'createPayment(idDon, phuongThuc)\nconfirmPayment(idThanhToan)\ngetPaymentStatus(idThanhToan)',
         'POST /payments\nPUT /payments/:id/confirm\nGET /payments/:id'],
        ['cancel.js', 'getCancelFee(idVe)\ncancelTickets(maDatCho, idVeList, lyDo)',
         'GET /cancel/fee/:idVe\nPOST /cancel'],
        ['exchange.js', 'checkExchangeable(idVe)\nexchangeTicket(idVeCu, newData)',
         'GET /exchange/check/:idVe\nPOST /exchange'],
    ],
    [3, 5.5, 6.5]
)

h2('5.4. Tiện ích Frontend (src/utils/)')
h3('authUtils.js – Quản lý JWT trong localStorage')
code_block("""const KLN_AUTH_KEY = 'KLN_AUTH'
getUser()           → JSON.parse(localStorage.getItem('KLN_AUTH'))
getToken()          → getUser()?.token
loginUser({token, user})  → localStorage.setItem('KLN_AUTH', JSON.stringify({token,...user}))
logoutUser()        → localStorage.removeItem('KLN_AUTH')
isLoggedIn()        → !!getToken()""")

h3('dateUtils.js')
bullet('formatDate(dateStr): chuyển YYYY-MM-DD hoặc ISO datetime thành DD/MM/YYYY.')
bullet('Trả "--/--/----" nếu chuỗi rỗng hoặc null.')

h2('5.5. Các trang (Pages) – src/pages/')
h3('Home (pages/home/)')
bullet('Hero.jsx: Banner chào mừng, nút CTA "Mua vé ngay" điều hướng tới /tim-ve.')
bullet('Search/Search.jsx: Form tìm kiếm chuyến tàu với 3 trường: ga đi, ga đến, ngày đi. Gọi getStations() để populate dropdown. Submit → navigate("/tim-ve", { state: { gaDi, gaDen, ngayChay } }).')
bullet('PopularRoutes/PopularRoutes.jsx: Hiển thị các tuyến phổ biến dưới dạng card (dữ liệu từ data/popularRoutes.js).')
bullet('TrainMap/TrainMap.jsx: Bản đồ tuyến đường tàu hỏa dạng hình ảnh trực quan.')

h3('TicketSearch (pages/ticketSearch/) – Trang tìm và chọn vé')
para('Trang trung tâm của quy trình đặt vé, gồm 3 giai đoạn:')
bullet('SearchForm/: Form cho phép sửa lại điều kiện tìm kiếm, submit lại gọi searchTrains().')
bullet('TrainSelection/: Hiển thị danh sách chuyến tàu từ API, mỗi item cho thấy giờ đi/đến, giá từ, số toa. Click → chuyển sang chọn ghế.')
bullet('SeatSelection/: Hiển thị sơ đồ ghế theo toa (getSeatMap()). Ghế màu xanh = trống, đỏ = đã đặt, vàng = đang giữ. Click ghế trống → thêm vào danh sách chọn → hiện nút "Tiếp tục".')
bullet('Submit: gọi holdSeats() → nhận sessionId → navigate("/checkout", { state: { selectedSeats, tripInfo, sessionId } }).')

h3('Checkout (pages/checkout/Checkout.jsx) – Điền thông tin và xác nhận')
para('Nhận state từ TicketSearch qua useLocation().state.')
bullet('Hiển thị tóm tắt chuyến tàu và ghế đã chọn.')
bullet('Form nhập thông tin hành khách (họ tên, ngày sinh, CCCD) + thông tin liên lạc (email, SĐT).')
bullet('Ô nhập mã khuyến mãi.')
bullet('Hiển thị bảng tính tiền: giá vé + phí dịch vụ - giảm giá = tổng thanh toán.')
bullet('Submit → gọi createBooking() → nhận { maDon, maDatCho, idDon, tongThanhToan } → navigate("/thanh-toan", { state: {...} }).')

h3('PaymentMethod (pages/payment/PaymentMethod.jsx)')
bullet('Nhận state: { idDon, tongThanhToan, maDatCho }.')
bullet('Hiển thị các phương thức thanh toán (QR Banking là chính, các hình thức khác UI only).')
bullet('Click "Thanh toán QR" → navigate("/thanh-toan/qr", { state }).')

h3('QRPayment (pages/payment/QRPayment.jsx)')
bullet('Gọi createPayment(idDon, "qr_bank") → nhận { idThanhToan, qrUrl, soTien, maDon }.')
bullet('Hiển thị ảnh QR từ VietQR URL với số tiền và nội dung chuyển khoản = maDon.')
bullet('Countdown 15 phút (setInterval, giảm mỗi giây, hiện MM:SS).')
bullet('Khi hết giờ → hiển thị thông báo hết hạn, disable nút xác nhận.')
bullet('Nút "Đã thanh toán" → gọi confirmPayment(idThanhToan) → navigate("/thanh-toan/thanh-cong", { state: { soHoaDon, maDatCho, tongThanhToan } }).')

h3('PaymentSuccess (pages/payment/PaymentSuccess.jsx)')
bullet('Trang xác nhận thanh toán thành công.')
bullet('Hiển thị số hóa đơn, mã đặt chỗ, số tiền đã thanh toán.')
bullet('Nút in vé → navigate("/in-ve") hoặc mở tab mới.')
bullet('Nút tra cứu đơn → navigate("/thong-tin-dat-cho").')

h3('BookingLookup (pages/bookingLookup/) – Tra cứu đặt vé')
bullet('BookingSearch/: Form nhập mã đặt chỗ + email hoặc SĐT → gọi lookupBooking().')
bullet('BookingHistory/: Nếu đã đăng nhập → gọi getBookingHistory() → hiển thị danh sách đơn.')
bullet('BookingResult/: Hiển thị chi tiết đơn (thông tin hành khách, chuyến tàu, vé, trạng thái).')

h3('CancelTicket (pages/cancelTicket/) – Hủy vé')
bullet('CancelSearch/: Nhập mã đặt chỗ + thông tin xác minh → gọi lookupBooking().')
bullet('CancelConfirm/: Chọn vé muốn hủy → gọi getCancelFee(idVe) → hiển thị phí hủy + tiền hoàn theo chính sách.')
bullet('Xác nhận → gọi cancelTickets(maDatCho, idVeList, lyDo).')
bullet('CancelSuccess/: Thông báo hủy thành công kèm số tiền được hoàn.')

h3('ExchangeTicket (pages/exchangeTicket/) – Đổi vé')
bullet('ExchangeSearch/: Nhập mã đặt chỗ → tìm đơn → chọn vé cần đổi.')
bullet('Gọi checkExchangeable(idVe) → hiển thị phí đổi, số giờ còn lại.')
bullet('ExchangeSelect/: Chọn chuyến mới, toa, ghế (tương tự SeatSelection).')
bullet('ExchangeConfirm/: So sánh vé cũ – vé mới, hiển thị phí đổi + chênh lệch giá.')
bullet('Xác nhận → gọi exchangeTicket(idVeCu, newTicketData).')

h3('TrainSchedule (pages/trainSchedule/) – Lịch chạy tàu')
bullet('Gọi getSchedule() khi mount → hiển thị bảng lịch tất cả tàu.')
bullet('SearchForm/: Filter theo ga đi, ga đến.')
bullet('TrainDetail/: Click vào lịch → gọi getTrainDetail() → hiển thị chi tiết: tất cả ga dừng, khoảng cách, giờ đến/đi, bảng giá theo loại ghế.')

h3('PrintTicket (pages/printTicket/PrintTicket.jsx)')
bullet('Không có Navbar, layout tối ưu cho in ấn.')
bullet('Nhận dữ liệu vé từ state hoặc query params.')
bullet('Hiển thị vé dạng card với: mã đặt chỗ, QR code (qrcode.react), thông tin hành khách, chuyến tàu, ghế.')
bullet('CSS print media query ẩn nút, chỉ hiện nội dung vé.')

h2('5.6. Components dùng chung')
h3('Navbar (component/Navbar/Navbar.jsx)')
bullet('Logo KLN Train + menu điều hướng chính.')
bullet('Đọc isLoggedIn() → hiển thị "Đăng nhập" hoặc tên người dùng + dropdown menu.')
bullet('Nút logout → logoutUser() + navigate("/").')

h3('TicketCard (components/TicketCard.jsx)')
bullet('Component dùng lại để hiển thị thẻ vé trong kết quả tra cứu, lịch sử đặt vé.')
bullet('Props: thông tin vé, hành khách, chuyến tàu, trạng thái.')

page_break()

# ═══════════════════════════════════════════════════════════════════
# VI. LUỒNG CHỨC NĂNG CHI TIẾT
# ═══════════════════════════════════════════════════════════════════
h1('VI. LUỒNG CHỨC NĂNG CHI TIẾT')
para('Phần này mô tả từng chức năng theo luồng dọc hoàn chỉnh: Frontend → API Client → Backend Route → Middleware → Controller → Service → Repository → Sequelize → SQL Server.')

h2('6.1. Đăng ký tài khoản')
add_table(
    ['Bước', 'Thành phần', 'Hành động / Truy vấn'],
    [
        ['1', 'Frontend (màn hình đăng ký)', 'User nhập email, mật khẩu, họ tên, SĐT → submit form'],
        ['2', 'api/auth.js → client.js', 'post("/auth/register", {email, matKhau, hoTen, soDienThoai})'],
        ['3', 'Express router (routes/auth.js)', 'POST /api/auth/register → express-validator kiểm tra email format, password ≥ 6 ký tự'],
        ['4', 'AuthController.register()', 'validationResult() → ném 400 nếu lỗi; gọi AuthService.register()'],
        ['5', 'AuthService.register()', 'TaiKhoanRepo.emailExists() → kiểm tra email; bcrypt.hash(matKhau, 10); TaiKhoan.create()'],
        ['6', 'SQL Server', 'INSERT INTO TaiKhoan (email, mat_khau, ho_ten, ...) VALUES (...)'],
        ['7', 'AuthService (tiếp)', 'signToken() → jwt.sign({id, email, role}, JWT_SECRET, {expiresIn:"24h"})'],
        ['8', 'AuthController → Response', 'created(res, {token, user}) → HTTP 201'],
        ['9', 'Frontend', 'loginUser({token, user}) → lưu vào localStorage["KLN_AUTH"]; navigate("/")'],
    ],
    [0.8, 4, 10.2]
)

h2('6.2. Đăng nhập')
add_table(
    ['Bước', 'Thành phần', 'Hành động / Truy vấn'],
    [
        ['1', 'Frontend', 'User nhập email + mật khẩu → submit'],
        ['2', 'api/auth.js', 'post("/auth/login", {email, matKhau})'],
        ['3', 'routes/auth.js', 'POST /api/auth/login → kiểm tra express-validator'],
        ['4', 'AuthController.login()', 'Gọi AuthService.login()'],
        ['5', 'AuthService.login()', 'TaiKhoanRepo.findByEmail(email) → SELECT * FROM TaiKhoan WHERE email=?'],
        ['6', 'AuthService (tiếp)', 'Kiểm tra trang_thai ≠ "khoa"; bcrypt.compare(matKhau, mat_khau)'],
        ['7', 'AuthService (tiếp)', 'signToken() → jwt.sign() → trả {token, user}'],
        ['8', 'Frontend', 'loginUser() → lưu localStorage; navigate về trang trước hoặc /'],
    ],
    [0.8, 4, 10.2]
)

h2('6.3. Tìm kiếm chuyến tàu')
add_table(
    ['Bước', 'Thành phần', 'Hành động / Truy vấn'],
    [
        ['1', 'Home / Search component', 'User chọn ga đi, ga đến, ngày → submit → navigate("/tim-ve", {state})'],
        ['2', 'TicketSearch mount', 'Đọc state từ useLocation() → gọi searchTrains(gaDi, gaDen, ngayChay)'],
        ['3', 'api/trains.js', 'get("/trains/search?gaDi=Hà Nội&gaDen=TP HCM&ngayChay=2026-06-01")'],
        ['4', 'routes/trains.js', 'GET /api/trains/search → không cần auth → TrainController.searchTrains()'],
        ['5', 'TrainController', 'Kiểm tra query params → gọi TrainService.searchTrains()'],
        ['6', 'TrainService.searchTrains()', 'GaTau.findOne WHERE ten_ga=? (x2) → lấy idGaDi, idGaDen'],
        ['7', 'TrainRepo.searchChuyen()', 'SELECT ChuyenTau JOIN LichChay (WHERE id_ga_di=?,id_ga_den=?) JOIN Tau JOIN GaTau WHERE ngay_chay=? AND trang_thai≠huy ORDER BY gio_khoi_hanh'],
        ['8', 'TrainService (tiếp)', 'Với mỗi chuyến: getCoachesByChuyen() → lấy danh sách toa; tinhGiaVe() → 2 raw SQL cho km + BieuGia.findAll'],
        ['9', 'Response', 'Trả mảng chuyến tàu với: idChuyen, giờ đi/đến, priceFrom, danh sách toa'],
        ['10', 'TicketSearch render', 'Hiển thị danh sách tàu. User click tàu → TrainSelection → chọn toa → SeatSelection'],
    ],
    [0.8, 4, 10.2]
)

h2('6.4. Xem sơ đồ ghế và chọn ghế')
add_table(
    ['Bước', 'Thành phần', 'Hành động / Truy vấn'],
    [
        ['1', 'SeatSelection mount', 'Sau khi chọn tàu và toa → gọi getSeatMap(idChuyen, soToa, idGaLen, idGaXuong)'],
        ['2', 'api/trains.js', 'get("/trains/:idChuyen/seats/:soToa?idGaLen=X&idGaXuong=Y")'],
        ['3', 'TrainController.getSeatMap()', 'Gọi TrainService.getSeatMap()'],
        ['4', 'TrainService.getSeatMap()', 'Gọi TrainRepo.getSeatMap()'],
        ['5', 'TrainRepo.getSeatMap()', 'Query 1: CauHinhGhe.findAll (WHERE id_loai_toa=?) include LoaiGhe → layout ghế\nQuery 2: Ve.findAll (WHERE id_chuyen=?, so_toa=?, trang_thai NOT IN (da_huy,da_doi)) → ghế đã đặt\nQuery 3: TamGiuGhe.findAll (WHERE id_chuyen=?, so_toa=?, trang_thai=dang_giu, het_han > GETUTCDATE()) → ghế đang giữ'],
        ['6', 'TrainService (tiếp)', 'Merge: empty/sold/held. Tính giá từng ghế qua tinhGiaVe() + priceCache'],
        ['7', 'SeatSelection render', 'Vẽ sơ đồ: ghế xanh (empty), đỏ (sold), vàng (held). Click ghế empty → thêm vào selectedSeats'],
        ['8', 'Nút Tiếp tục', 'gọi holdSeats({ trips:[{idChuyen, passengerSeats:[{soToaThuTu,seatNumber}]}] })'],
        ['9', 'BookingService.holdSeatsForCheckout()', 'crypto.randomBytes(8).hex() → sessionId; checkSeatsAvailable(); BookingRepo.holdSeats() → INSERT TamGiuGhe'],
        ['10', 'Frontend', 'navigate("/checkout", { state: { selectedSeats, sessionId, hetHan } })'],
    ],
    [0.8, 4, 10.2]
)

h2('6.5. Checkout – Tạo đơn đặt vé')
add_table(
    ['Bước', 'Thành phần', 'Hành động / Truy vấn'],
    [
        ['1', 'Checkout.jsx mount', 'Đọc state: selectedSeats, sessionId, tripInfo'],
        ['2', 'User nhập thông tin', 'Họ tên HK, ngày sinh, CCCD cho từng hành khách; email, SĐT liên lạc; mã KM (tùy chọn)'],
        ['3', 'Submit form', 'gọi createBooking({ trips, passengers, contactInfo, maKhuyenMai })'],
        ['4', 'api/bookings.js', 'post("/bookings", body)'],
        ['5', 'routes/bookings.js', 'POST /api/bookings → optionalAuth → BookingController.createBooking()'],
        ['6', 'BookingController', 'req.body + req.user?.id → BookingService.createBooking()'],
        ['7', 'BookingService – BEGIN TRANSACTION', ''],
        ['7a', '  Validate KM', 'KhuyenMai.findOne WHERE ma_khuyen_mai=? → kiểm tra ngày + số lượt'],
        ['7b', '  Tính tiền', 'tongTienVe = Σ seatPrice; phiDV = 20k/40k; tienGiam từ KM'],
        ['7c', '  Kiểm tra ghế', 'BookingRepo.checkSeatsAvailable() – Ve.count + TamGiuGhe.count (exclude sessionId)'],
        ['7d', '  Sinh mã đơn', 'genBookingCode() + vòng lặp unique check; genOrderCode()'],
        ['7e', '  Tạo DonDatVe', 'INSERT DonDatVe (ma_don, ma_dat_cho, tong_tien, trang_thai=cho_thanh_toan, het_han=now+15m)'],
        ['7f', '  Tạo HK + Ve', 'HanhKhach.findOne (ten+ngaySinh) or create; Ve.create() cho mỗi HK'],
        ['7g', '  Link TamGiuGhe', 'TamGiuGhe.update SET id_don_dat_ve=? WHERE session_id=? AND id_chuyen=?'],
        ['7h', '  Cập nhật KM', 'KhuyenMai.increment("da_dung", {by:1})'],
        ['7i', '  COMMIT', ''],
        ['8', 'Response → Frontend', 'created(res, {maDon, maDatCho, idDon, tongThanhToan}) → navigate("/thanh-toan")'],
    ],
    [0.8, 4.5, 9.7]
)

h2('6.6. Thanh toán QR VietQR')
add_table(
    ['Bước', 'Thành phần', 'Hành động / Truy vấn'],
    [
        ['1', 'PaymentMethod', 'User chọn "Chuyển khoản QR" → navigate("/thanh-toan/qr", { state: { idDon, tongThanhToan, maDatCho } })'],
        ['2', 'QRPayment mount', 'gọi createPayment(idDon, "qr_bank")'],
        ['3', 'api/payments.js', 'post("/payments", { idDonDatVe: idDon, phuongThuc: "qr_bank" })'],
        ['4', 'PaymentController.createPayment()', 'Gọi PaymentService.createPayment()'],
        ['5', 'PaymentService.createPayment()', 'DonDatVe.findByPk → kiểm tra trang_thai=cho_thanh_toan + chưa hết 15 phút\ngenTransactionCode() → UUID 20 ký tự\nqrUrl = "https://img.vietqr.io/image/BIDV-...?amount=X&addInfo=maDon"\nThanhToan.create()'],
        ['6', 'Response', '{ idThanhToan, maGiaoDich, qrUrl, soTien, maDon }'],
        ['7', 'QRPayment render', 'Hiển thị <img src={qrUrl} />, countdown 15 phút (setInterval mỗi giây)'],
        ['8', 'User quét QR + chuyển khoản', '(Thao tác ngoài hệ thống – dùng app ngân hàng)'],
        ['9', 'User nhấn "Đã thanh toán"', 'gọi confirmPayment(idThanhToan)'],
        ['10', 'PaymentService.confirmPayment() – BEGIN TRANSACTION', ''],
        ['10a', '  Cập nhật ThanhToan', 'UPDATE ThanhToan SET trang_thai=thanh_cong, thoi_gian_thanh_toan=NOW'],
        ['10b', '  Cập nhật DonDatVe', 'UPDATE DonDatVe SET trang_thai=da_thanh_toan'],
        ['10c', '  Xác nhận Ve', 'UPDATE Ve SET trang_thai=da_xac_nhan WHERE id_don=? AND trang_thai=cho_xac_nhan'],
        ['10d', '  Chốt TamGiuGhe', 'UPDATE TamGiuGhe SET trang_thai=da_dat WHERE id_don=?'],
        ['10e', '  Tạo HoaDon', 'genInvoiceNumber() → INSERT HoaDon (so_hoa_don, tong_tien, tien_giam, ...)'],
        ['10f', '  COMMIT', ''],
        ['11', 'Response → Frontend', '{ soHoaDon, maDatCho, tongThanhToan } → navigate("/thanh-toan/thanh-cong")'],
    ],
    [0.8, 4.5, 9.7]
)

h2('6.7. Tra cứu đơn đặt vé')
add_table(
    ['Bước', 'Thành phần', 'Hành động / Truy vấn'],
    [
        ['1', 'BookingLookup / BookingSearch', 'User nhập mã đặt chỗ (6 ký tự) + email hoặc SĐT'],
        ['2', 'api/bookings.js', 'post("/bookings/lookup", { maDatCho, email, phone })'],
        ['3', 'BookingController.lookupBooking()', 'Gọi BookingService.lookupBooking()'],
        ['4', 'BookingService.lookupBooking()', 'BookingRepo.findByMaDatCho(maDatCho.toUpperCase()) → SELECT DonDatVe JOIN Ve JOIN HanhKhach JOIN ChuyenTau...'],
        ['5', 'BookingService (tiếp)', 'So khớp email (case-insensitive) HOẶC số điện thoại (7 số cuối). Trả null nếu không khớp.'],
        ['6', 'BookingService.formatDon()', 'Chuyển object Sequelize → flat JSON camelCase cho frontend'],
        ['7', 'Frontend render', 'Hiển thị thông tin đơn: mã đơn, trạng thái, danh sách vé, hành khách, chuyến tàu'],
    ],
    [0.8, 4, 10.2]
)

h2('6.8. Hủy vé')
add_table(
    ['Bước', 'Thành phần', 'Hành động / Truy vấn'],
    [
        ['1', 'CancelTicket / CancelSearch', 'Tra cứu đơn → chọn vé muốn hủy'],
        ['2', 'CancelConfirm mount', 'gọi getCancelFee(idVe) cho từng vé chọn'],
        ['3', 'api/cancel.js', 'get("/cancel/fee/:idVe")'],
        ['4', 'CancelController.getCancelFee()', 'Gọi CancelService.tinhPhiHuy(idVe)'],
        ['5', 'CancelService.tinhPhiHuy()', 'Ve.findByPk include ChuyenTau → LichChay → tính gioConLai\nChinhSachHuy.findOne (WHERE gio_truoc_gio_chay ≥ floor(gioConLai) ORDER ASC) → phiHuyPct\nTrả { giaVe, phiHuyPct, phiHuy, tienHoan, gioConLai, canCancel }'],
        ['6', 'Frontend', 'Hiển thị bảng: giá vé, phí hủy %, tiền hoàn. User xác nhận'],
        ['7', 'api/cancel.js', 'post("/cancel", { maDatCho, idVeList, lyDo })'],
        ['8', 'CancelService.cancelTickets() – TRANSACTION', 'DonDatVe.findOne include Ve + ThanhToan\nVới mỗi vé: ve.update(da_huy) + HoanTien.create(tien_goc, phi_huy, tien_hoan)\nNếu all hủy: don.update(da_huy)\nTrả { tongTienHoan, soVeHuy }'],
        ['9', 'Frontend', 'CancelSuccess: hiển thị "Hủy thành công, hoàn X đồng"'],
    ],
    [0.8, 4, 10.2]
)

h2('6.9. Đổi vé')
add_table(
    ['Bước', 'Thành phần', 'Hành động / Truy vấn'],
    [
        ['1', 'ExchangeSearch', 'Tra cứu đơn → chọn vé cần đổi'],
        ['2', 'checkExchangeable(idVe)', 'GET /exchange/check/:idVe → trang_thai=da_xac_nhan, gioConLai≥24h; tính phiDoi = max(5%, 20k)'],
        ['3', 'ExchangeSelect', 'Chọn chuyến tàu mới, toa, ghế (tương tự quy trình chọn ghế)'],
        ['4', 'ExchangeConfirm', 'Tính chênh lệch giá (giaVeMoi - giaVeCu, nếu >0); tongPhaiTra = phiDoi + chenhLech'],
        ['5', 'exchangeTicket(idVeCu, newData)', 'POST /exchange → TRANSACTION: checkSeatsAvailable() → Ve.create(vé mới) → DoiVe.create() → veCu.update(da_doi)'],
        ['6', 'Frontend', 'Hiển thị thông tin vé mới, số tiền phải thanh toán thêm (nếu có)'],
    ],
    [0.8, 4, 10.2]
)

h2('6.10. Xem lịch chạy tàu')
add_table(
    ['Bước', 'Thành phần', 'Hành động / Truy vấn'],
    [
        ['1', 'TrainSchedule mount', 'gọi getSchedule()'],
        ['2', 'api/trains.js', 'get("/trains/schedule")'],
        ['3', 'TrainController.getSchedule()', 'Gọi TrainService.getSchedule()'],
        ['4', 'TrainService.getSchedule()', 'LichChay.findAll include Tau, GaTau(GaDi/GaDen) ORDER BY gio_khoi_hanh'],
        ['5', 'User click chi tiết', 'gọi getTrainDetail(idLichChay, idGaLen, idGaXuong, ngayChay)'],
        ['6', 'TrainService.getTrainRouteDetail()', 'LichChay.findByPk include Tau → CauHinhToa → LoaiToa → LoaiGhe\ngetLichTrinh() → SELECT LichTrinhChuyen JOIN GaTau\ntinhGiaVe() cho từng loại ghế'],
        ['7', 'Frontend', 'Hiển thị: tất cả ga dừng (km, giờ đến/đi), bảng giá theo loại ghế'],
    ],
    [0.8, 4, 10.2]
)

page_break()

# ═══════════════════════════════════════════════════════════════════
# VII. BẢO MẬT VÀ XỬ LÝ LỖI
# ═══════════════════════════════════════════════════════════════════
h1('VII. BẢO MẬT VÀ XỬ LÝ LỖI')

h2('7.1. Bảo mật')
add_table(
    ['Cơ chế', 'Thực hiện'],
    [
        ['Mã hóa mật khẩu', 'bcryptjs với 10 salt rounds. Mật khẩu không bao giờ lưu dạng plaintext.'],
        ['Xác thực JWT', 'Token 24h, ký bằng JWT_SECRET từ .env. Middleware authenticate verify mỗi request.'],
        ['CORS', 'Chỉ cho phép origin = FRONTEND_URL (http://localhost:5173), credentials: true.'],
        ['SQL Injection', 'Dùng Sequelize ORM với parameterized queries. Raw SQL dùng replacements (không nối chuỗi).'],
        ['Input Validation', 'express-validator trên tất cả endpoint auth. Kiểm tra email format, độ dài password.'],
        ['Phân quyền', 'requireAdmin middleware kiểm tra vai_tro = "quan_tri" cho endpoint admin.'],
        ['Giới hạn payload', 'express.json({ limit: "10mb" }) ngăn chặn oversized request.'],
    ],
    [4, 11]
)

h2('7.2. Xử lý lỗi')
add_table(
    ['Tình huống', 'Cơ chế xử lý'],
    [
        ['Route không tồn tại', 'notFoundHandler → trả HTTP 404 với message "Endpoint không tồn tại"'],
        ['JWT hết hạn / sai', 'authenticate() → unauthorized(res, "Token không hợp lệ") → HTTP 401'],
        ['Validation thất bại', 'validationResult() trong Controller → badRequest(res, errors) → HTTP 400'],
        ['Ghế đã được đặt', 'checkSeatsAvailable() = false → throw { status: 409, message: "Ghế đã được đặt" }'],
        ['Sequelize UniqueConstraint', 'errorHandler catch → HTTP 409 "Dữ liệu đã tồn tại"'],
        ['Sequelize ValidationError', 'errorHandler catch → HTTP 400 kèm danh sách lỗi từng trường'],
        ['Lỗi không xác định', 'errorHandler → HTTP 500 "Lỗi máy chủ nội bộ"'],
        ['Frontend – network error', 'Axios response interceptor → Error("Không thể kết nối máy chủ")'],
        ['Frontend – timeout 20s', 'Axios response interceptor → Error("Yêu cầu quá lâu, vui lòng thử lại")'],
    ],
    [5, 10]
)

h2('7.3. Sequelize Transaction')
para('Các nghiệp vụ đa bảng sử dụng sequelize.transaction() để đảm bảo tính toàn vẹn dữ liệu (ACID):')
add_table(
    ['Nghiệp vụ', 'Bảng liên quan trong transaction'],
    [
        ['createBooking()', 'DonDatVe + Ve + HanhKhach + TamGiuGhe + KhuyenMai'],
        ['confirmPayment()', 'ThanhToan + DonDatVe + Ve + TamGiuGhe + HoaDon'],
        ['cancelTickets()', 'Ve + HoanTien + DonDatVe'],
        ['exchangeTicket()', 'Ve (mới) + Ve (cũ, update) + DoiVe'],
    ],
    [5, 10]
)

page_break()

# ═══════════════════════════════════════════════════════════════════
# VIII. CẤU HÌNH MÔI TRƯỜNG
# ═══════════════════════════════════════════════════════════════════
h1('VIII. CẤU HÌNH MÔI TRƯỜNG')

h2('8.1. Backend – backend/.env')
code_block("""# Kết nối SQL Server (Named Instance)
DB_SERVER=DESKTOP-5DF60PC\\SQLEXPRESS
DB_PORT=1433
DB_NAME=KLNTrain
DB_USER=TRONGLINH
DB_PASSWORD=<password>

# JWT
JWT_SECRET=klntrain_super_secret_key_change_in_production
JWT_EXPIRES_IN=24h

# Server
PORT=8000
NODE_ENV=development
FRONTEND_URL=http://localhost:5173""")

h2('8.2. Frontend – frontend/.env')
code_block("""VITE_API_URL=http://localhost:8000/api""")

h2('8.3. Hướng dẫn chạy dự án')
add_table(
    ['Bước', 'Lệnh', 'Ghi chú'],
    [
        ['1. Cài backend', 'cd backend && npm install', 'Cài express, sequelize, mssql, bcryptjs, jwt…'],
        ['2. Cài frontend', 'cd frontend && npm install', 'Cài react, vite, tailwindcss, axios…'],
        ['3. Restore DB', 'Restore file kln_train.sql vào SQL Server', 'Đặt tên DB = KLNTrain'],
        ['4. Cấu hình .env', 'Sửa backend/.env', 'Nhập đúng DB_SERVER, DB_USER, DB_PASSWORD'],
        ['5. Chạy backend', 'cd backend && npm run dev', 'nodemon app.js → http://localhost:8000'],
        ['6. Chạy frontend', 'cd frontend && npm run dev', 'Vite HMR → http://localhost:5173'],
        ['7. Kiểm tra', 'GET http://localhost:8000/api/health', 'Trả { status: "OK" }'],
    ],
    [1.5, 5.5, 8]
)

page_break()

# ═══════════════════════════════════════════════════════════════════
# IX. KẾT LUẬN
# ═══════════════════════════════════════════════════════════════════
h1('IX. KẾT LUẬN')
para('Hệ thống KLN Train là một ứng dụng đặt vé tàu hỏa trực tuyến được xây dựng hoàn chỉnh theo kiến trúc Full-stack hiện đại, gồm 3 thành phần chính:', bold=False)
bullet('Backend Node.js/Express với kiến trúc phân tầng rõ ràng (Route → Middleware → Controller → Service → Repository → Model), đảm bảo tách biệt trách nhiệm và dễ mở rộng.')
bullet('Frontend React 19 SPA với Vite, áp dụng React Router cho điều hướng và Axios với interceptors cho quản lý JWT và xử lý lỗi tập trung.')
bullet('Cơ sở dữ liệu SQL Server với 25+ bảng, phủ đầy đủ các nghiệp vụ: quản lý tàu, lịch chạy, đặt vé, thanh toán, hủy/đổi vé.')

h2('Các điểm kỹ thuật nổi bật')
add_table(
    ['Tính năng', 'Kỹ thuật áp dụng'],
    [
        ['Ngăn chặn đặt trùng ghế', 'TamGiuGhe + checkSeatsAvailable() với GETUTCDATE() trong SQL Server'],
        ['Tính giá vé động', 'Raw SQL lấy km + BieuGia ORM + công thức làm tròn 1.000đ'],
        ['Giao dịch toàn vẹn', 'sequelize.transaction() cho mọi nghiệp vụ đa bảng'],
        ['Thanh toán QR', 'VietQR API tạo ảnh QR với số tài khoản BIDV + số tiền + nội dung'],
        ['Chính sách hủy linh hoạt', 'DB-driven ChinhSachHuy, admin điều chỉnh không cần deploy lại'],
        ['Bảo mật JWT', 'Stateless, 24h expiry, middleware xác thực mọi request cần auth'],
        ['Audit Trail', 'AuditLog table ghi lại mọi thay đổi dữ liệu quan trọng'],
    ],
    [5, 10]
)

para('Dự án đáp ứng đầy đủ các yêu cầu nghiệp vụ của một hệ thống đặt vé tàu hỏa thực tế, từ tìm kiếm chuyến, chọn ghế, đặt vé, thanh toán đến hủy và đổi vé, đồng thời áp dụng các kỹ thuật bảo mật và xử lý đồng thời phù hợp.')

# ── Save ──────────────────────────────────────────────────────────
output_path = r'd:\Nam3\Ki2_Nam3\CNLTTH\DuAnTauHoa\BaoCao_DuAnTauHoa_KLNTrain.docx'
doc.save(output_path)
print(f'Done: {output_path}')
