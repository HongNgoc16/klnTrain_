#!/usr/bin/env python3
"""
generate_sql_test.py
Generates KLN_Train_SQL_Test.docx — a professional DOCX document with
comprehensive SQL test queries for the KLN Train ticketing system.
"""

import subprocess, sys

# ── Ensure python-docx is available ──────────────────────────────────────────
try:
    from docx import Document
except ImportError:
    print("python-docx not found — installing …")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"d:\Nam3\Ki2_Nam3\CNLTTH\DuAnTauHoa\KLN_Train_SQL_Test.docx"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_shading(paragraph, fill_hex: str):
    """Apply solid background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def set_run_font(run, name: str, size_pt: float, bold=False, color_rgb=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)


def add_normal(doc, text: str, bold=False, color_rgb=None, size_pt=11):
    """Add a normal Calibri paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    set_run_font(run, 'Calibri', size_pt, bold=bold, color_rgb=color_rgb)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_code_block(doc, code: str):
    """Add a SQL code block: Courier New 9pt on light-gray (#F5F5F5) background."""
    lines = code.strip('\n').split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        set_paragraph_shading(p, 'F5F5F5')
        run = p.add_run(line if line else ' ')
        set_run_font(run, 'Courier New', 9)
    # Small gap after code block
    gap = doc.add_paragraph()
    gap.style = doc.styles['Normal']
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(6)


def add_h1(doc, text: str):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h2(doc, text: str):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_h3(doc, text: str):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x69, 0x5C)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_note(doc, text: str):
    """Add a bold red note paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    set_run_font(run, 'Calibri', 11, bold=True, color_rgb=(0xC0, 0x00, 0x00))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_page_break(doc):
    doc.add_page_break()


def add_toc_entry(doc, number: str, title: str):
    """Simulate a TOC line."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    r1 = p.add_run(f"{number}  ")
    set_run_font(r1, 'Calibri', 11, bold=True)
    r2 = p.add_run(title)
    set_run_font(r2, 'Calibri', 11)
    p.paragraph_format.space_after = Pt(2)


# ── Build the document ────────────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.5)

    # ═══════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════════════════
    for _ in range(4):
        blank = doc.add_paragraph()
        blank.paragraph_format.space_after = Pt(0)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("KLN TRAIN")
    set_run_font(r, 'Calibri', 28, bold=True, color_rgb=(0x1F, 0x49, 0x7D))

    sub1_p = doc.add_paragraph()
    sub1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub1_p.add_run("TÀI LIỆU KIỂM TRA CƠ SỞ DỮ LIỆU SQL SERVER")
    set_run_font(r, 'Calibri', 18, bold=True, color_rgb=(0x2E, 0x74, 0xB5))

    doc.add_paragraph()

    sub2_p = doc.add_paragraph()
    sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2_p.add_run(
        "Hướng dẫn kiểm tra tính đúng đắn của dữ liệu và truy vấn\n"
        "cho tất cả chức năng hệ thống"
    )
    set_run_font(r, 'Calibri', 13, color_rgb=(0x40, 0x40, 0x40))

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date_p.add_run("2026")
    set_run_font(r, 'Calibri', 12, color_rgb=(0x60, 0x60, 0x60))

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════════════════
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = toc_title.add_run("MỤC LỤC")
    set_run_font(r, 'Calibri', 16, bold=True, color_rgb=(0x1F, 0x49, 0x7D))
    toc_title.paragraph_format.space_after = Pt(14)

    toc_entries = [
        ("1",   "KIỂM TRA CẤU TRÚC DỮ LIỆU CƠ BẢN"),
        ("1.1", "Kiểm tra danh sách ga tàu"),
        ("1.2", "Kiểm tra tàu và lịch chạy"),
        ("1.3", "Kiểm tra chuyến tàu theo ngày (ChuyenTau)"),
        ("1.4", "Kiểm tra cấu hình toa (CauHinhToa & LoaiToa)"),
        ("1.5", "Kiểm tra cấu hình ghế (CauHinhGhe & LoaiGhe)"),
        ("2",   "KIỂM TRA CHỨC NĂNG TÌM CHUYẾN TÀU"),
        ("2.1", "Kiểm tra lịch trình và khoảng cách km (LichTrinhChuyen)"),
        ("2.2", "Kiểm tra biểu giá (BieuGia)"),
        ("2.3", "Mô phỏng toàn bộ luồng tìm chuyến"),
        ("3",   "KIỂM TRA CHỨC NĂNG ĐẶT VÉ"),
        ("3.1", "Kiểm tra đơn đặt vé"),
        ("3.2", "Kiểm tra giữ ghế tạm thời (TamGiuGhe)"),
        ("4",   "KIỂM TRA CHỨC NĂNG TRA CỨU ĐẶT CHỖ"),
        ("4.1", "Tra cứu đơn theo mã đặt chỗ + email + SĐT"),
        ("5",   "KIỂM TRA CHỨC NĂNG HỦY VÉ"),
        ("5.1", "Kiểm tra chính sách hủy (ChinhSachHuy)"),
        ("6",   "KIỂM TRA CHỨC NĂNG ĐỔI VÉ"),
        ("6.1", "Kiểm tra điều kiện và kết quả đổi vé"),
        ("7",   "KIỂM TRA KHUYẾN MÃI"),
        ("8",   "CHÈN DỮ LIỆU MẪU ĐẦY ĐỦ"),
        ("8.1", "Dữ liệu mẫu: BieuGia"),
        ("8.2", "Dữ liệu mẫu: ChinhSachHuy"),
        ("8.3", "Dữ liệu mẫu: ChuyenTau cho các ngày test"),
        ("8.4", "Kiểm tra LichTrinhChuyen và thêm nếu thiếu"),
        ("9",   "BÁO CÁO TỔNG QUAN HỆ THỐNG"),
    ]
    for num, title in toc_entries:
        indent = "    " if "." in num else ""
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        r = p.add_run(f"{indent}{num}   {title}")
        bold = "." not in num
        set_run_font(r, 'Calibri', 11, bold=bold)
        p.paragraph_format.space_after = Pt(3)

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # PHẦN 1
    # ═══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "PHẦN 1: KIỂM TRA CẤU TRÚC DỮ LIỆU CƠ BẢN")

    # 1.1
    add_h2(doc, "1.1 Kiểm tra danh sách ga tàu")
    add_normal(doc, "Mục đích: Xác nhận bảng GaTau có dữ liệu và tên ga khớp với những gì frontend gửi lên.")
    add_code_block(doc, """\
-- 1.1.1 Xem tất cả ga đang hoạt động (phải có ít nhất 5-10 ga)
SELECT id_ga, ten_ga, ma_ga_viet_tat, do_uu_tien, trang_thai
FROM GaTau
WHERE trang_thai = 'hoat_dong'
ORDER BY do_uu_tien DESC;

-- 1.1.2 Kiểm tra các ga mà frontend thường gửi (tên phải khớp chính xác)
-- Thay 'Hà Nội', 'Sài Gòn' bằng tên ga thực tế trong DB của bạn
SELECT id_ga, ten_ga, trang_thai
FROM GaTau
WHERE ten_ga IN (N'Hà Nội', N'Sài Gòn', N'Đà Nẵng', N'Nha Trang', N'Huế')
ORDER BY ten_ga;

-- 1.1.3 Tìm ga theo từ khóa (để debug khi tên không khớp chính xác)
SELECT id_ga, ten_ga, ma_ga_viet_tat, trang_thai
FROM GaTau
WHERE ten_ga LIKE N'%Hà Nội%'
   OR ten_ga LIKE N'%Ha Noi%'
   OR ten_ga LIKE N'%Sài Gòn%'
   OR ten_ga LIKE N'%Sai Gon%';""")

    # 1.2
    add_h2(doc, "1.2 Kiểm tra tàu và lịch chạy")
    add_normal(doc, "Mục đích: Đảm bảo bảng Tau, LichChay có dữ liệu và liên kết đúng.")
    add_code_block(doc, """\
-- 1.2.1 Xem tất cả tàu trong hệ thống
SELECT id_tau, so_hieu, ten_tau, so_toa, trang_thai
FROM Tau
ORDER BY so_hieu;

-- 1.2.2 Xem lịch chạy và liên kết ga đi/ga đến
SELECT
    lc.id_lich_chay,
    t.so_hieu AS ma_tau,
    t.ten_tau,
    gd.ten_ga AS ga_di,
    gn.ten_ga AS ga_den,
    lc.gio_khoi_hanh,
    lc.gio_du_kien_den,
    lc.trang_thai
FROM LichChay lc
JOIN Tau t ON lc.id_tau = t.id_tau
JOIN GaTau gd ON lc.id_ga_di = gd.id_ga
JOIN GaTau gn ON lc.id_ga_den = gn.id_ga
ORDER BY lc.gio_khoi_hanh;

-- 1.2.3 Kiểm tra lịch chạy giữa 2 ga cụ thể (thay id_ga cho đúng)
DECLARE @idGaDi   INT = 1;   -- Thay bằng id_ga của ga đi
DECLARE @idGaDen  INT = 2;   -- Thay bằng id_ga của ga đến
SELECT
    lc.id_lich_chay,
    t.so_hieu,
    lc.gio_khoi_hanh,
    lc.gio_du_kien_den
FROM LichChay lc
JOIN Tau t ON lc.id_tau = t.id_tau
WHERE lc.id_ga_di = @idGaDi AND lc.id_ga_den = @idGaDen
  AND lc.trang_thai != 'huy';""")

    # 1.3
    add_h2(doc, "1.3 Kiểm tra chuyến tàu theo ngày (ChuyenTau)")
    add_note(doc,
        "ĐÂY LÀ ĐIỂM HAY XẢY RA LỖI: Nếu không có bản ghi trong ChuyenTau, "
        "API tìm chuyến sẽ trả về rỗng dù LichChay đã đúng.")
    add_normal(doc,
        "Đây là bảng then chốt: mỗi chuyến là một ngày cụ thể của một lịch chạy. "
        "Nếu không có bản ghi trong ChuyenTau, API tìm chuyến sẽ trả về rỗng dù LichChay đã đúng.")
    add_code_block(doc, """\
-- 1.3.1 Xem tất cả chuyến tàu (kiểm tra có dữ liệu không)
SELECT
    ct.id_chuyen,
    ct.ngay_chay,
    ct.trang_thai,
    lc.id_lich_chay,
    t.so_hieu AS ma_tau,
    gd.ten_ga AS ga_di,
    gn.ten_ga AS ga_den,
    lc.gio_khoi_hanh
FROM ChuyenTau ct
JOIN LichChay lc ON ct.id_lich_chay = lc.id_lich_chay
JOIN Tau t ON lc.id_tau = t.id_tau
JOIN GaTau gd ON lc.id_ga_di = gd.id_ga
JOIN GaTau gn ON lc.id_ga_den = gn.id_ga
ORDER BY ct.ngay_chay DESC, lc.gio_khoi_hanh;

-- 1.3.2 Tìm chuyến theo đúng ngày mà bạn đang test (QUAN TRỌNG)
DECLARE @ngayTest DATE = CAST(GETDATE() AS DATE); -- Hoặc thay bằng ngày cụ thể
SELECT
    ct.id_chuyen,
    ct.ngay_chay,
    t.so_hieu,
    gd.ten_ga AS ga_di,
    gn.ten_ga AS ga_den,
    lc.gio_khoi_hanh,
    ct.trang_thai
FROM ChuyenTau ct
JOIN LichChay lc ON ct.id_lich_chay = lc.id_lich_chay
JOIN Tau t ON lc.id_tau = t.id_tau
JOIN GaTau gd ON lc.id_ga_di = gd.id_ga
JOIN GaTau gn ON lc.id_ga_den = gn.id_ga
WHERE ct.ngay_chay = @ngayTest
  AND ct.trang_thai != 'huy'
ORDER BY lc.gio_khoi_hanh;

-- 1.3.3 Thống kê chuyến tàu theo tháng (để biết dữ liệu có trong khoảng nào)
SELECT
    YEAR(ngay_chay) AS nam,
    MONTH(ngay_chay) AS thang,
    COUNT(*) AS so_chuyen
FROM ChuyenTau
WHERE trang_thai != 'huy'
GROUP BY YEAR(ngay_chay), MONTH(ngay_chay)
ORDER BY nam DESC, thang DESC;

-- 1.3.4 Nếu thiếu chuyến, INSERT mẫu cho ngày hiện tại
-- (Chạy sau khi đã biết id_lich_chay từ query 1.2.2)
/*
DECLARE @idLichChay INT = 1; -- Thay bằng id_lich_chay thực tế
DECLARE @ngayChay DATE = '2026-06-01';
INSERT INTO ChuyenTau (id_lich_chay, ngay_chay, trang_thai)
VALUES (@idLichChay, @ngayChay, 'dang_chay');
*/""")

    # 1.4
    add_h2(doc, "1.4 Kiểm tra cấu hình toa (CauHinhToa & LoaiToa)")
    add_code_block(doc, """\
-- 1.4.1 Xem loại toa hiện có
SELECT id_loai_toa, ma_loai_toa, ten_loai_toa, so_cho_toi_da, trang_thai
FROM LoaiToa
ORDER BY ma_loai_toa;
-- Kết quả mong đợi: NMCLC (56 chỗ), GN6AC (60 chỗ), GN4AC (40 chỗ)

-- 1.4.2 Xem cấu hình toa của từng tàu
SELECT
    ct.id_cau_hinh,
    t.so_hieu AS ma_tau,
    ct.so_toa_thu_tu,
    lt.ma_loai_toa,
    lt.ten_loai_toa,
    lt.so_cho_toi_da
FROM CauHinhToa ct
JOIN Tau t ON ct.id_tau = t.id_tau
JOIN LoaiToa lt ON ct.id_loai_toa = lt.id_loai_toa
ORDER BY t.so_hieu, ct.so_toa_thu_tu;

-- 1.4.3 Kiểm tra tàu cụ thể có đủ toa không (thay id_tau)
DECLARE @idTau INT = 1;
SELECT
    ct.so_toa_thu_tu,
    lt.ma_loai_toa,
    lt.ten_loai_toa
FROM CauHinhToa ct
JOIN LoaiToa lt ON ct.id_loai_toa = lt.id_loai_toa
WHERE ct.id_tau = @idTau
ORDER BY ct.so_toa_thu_tu;""")

    # 1.5
    add_h2(doc, "1.5 Kiểm tra cấu hình ghế (CauHinhGhe & LoaiGhe)")
    add_code_block(doc, """\
-- 1.5.1 Xem loại ghế và hệ số giá
SELECT id_loai_ghe, ten_loai_ghe, ma_loai_ghe, he_so_ghe, mo_ta
FROM LoaiGhe
ORDER BY he_so_ghe;

-- 1.5.2 Kiểm tra số lượng ghế mỗi loại toa (phải đủ theo so_cho_toi_da)
SELECT
    lt.ma_loai_toa,
    lt.ten_loai_toa,
    lt.so_cho_toi_da,
    COUNT(cg.id_cau_hinh_ghe) AS so_ghe_cau_hinh
FROM LoaiToa lt
LEFT JOIN CauHinhGhe cg ON lt.id_loai_toa = cg.id_loai_toa
GROUP BY lt.id_loai_toa, lt.ma_loai_toa, lt.ten_loai_toa, lt.so_cho_toi_da
ORDER BY lt.ma_loai_toa;

-- 1.5.3 Xem chi tiết ghế của một loại toa (ví dụ NMCLC)
DECLARE @idLoaiToa INT = 1; -- Thay bằng id của NMCLC
SELECT
    cg.so_ghe_trong_toa,
    lg.ten_loai_ghe,
    lg.he_so_ghe,
    cg.vi_tri,
    cg.tang,
    cg.khoang_so,
    cg.ben
FROM CauHinhGhe cg
JOIN LoaiGhe lg ON cg.id_loai_ghe = lg.id_loai_ghe
WHERE cg.id_loai_toa = @idLoaiToa
ORDER BY cg.so_ghe_trong_toa;""")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # PHẦN 2
    # ═══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "PHẦN 2: KIỂM TRA CHỨC NĂNG TÌM CHUYẾN TÀU")

    # 2.1
    add_h2(doc, "2.1 Kiểm tra lịch trình và khoảng cách km (LichTrinhChuyen)")
    add_note(doc,
        "ĐÂY LÀ ĐIỂM HAY XẢY RA LỖI: Nếu thiếu bản ghi trong LichTrinhChuyen, "
        "hàm tinhGiaVe() sẽ trả về null dẫn đến priceFrom = null và frontend hiển thị giá 0.")
    add_code_block(doc, """\
-- 2.1.1 Xem toàn bộ lịch trình (các điểm dừng và km)
SELECT
    lt.id_lich_trinh,
    lc.id_lich_chay,
    t.so_hieu AS ma_tau,
    g.ten_ga,
    lt.thu_tu_dung,
    lt.khoang_cach_km,
    lt.gio_du_kien_den,
    lt.gio_du_kien_di
FROM LichTrinhChuyen lt
JOIN LichChay lc ON lt.id_lich_chay = lc.id_lich_chay
JOIN Tau t ON lc.id_tau = t.id_tau
JOIN GaTau g ON lt.id_ga = g.id_ga
ORDER BY t.so_hieu, lt.thu_tu_dung;

-- 2.1.2 Kiểm tra km của 2 ga trong một lịch chạy cụ thể
-- (Đây là query mà backend dùng để tính km cho tinhGiaVe)
DECLARE @idLichChay INT = 1;
DECLARE @idGaLen    INT = 1; -- id_ga của ga đi
DECLARE @idGaXuong  INT = 2; -- id_ga của ga đến

SELECT
    g.ten_ga,
    lt.khoang_cach_km,
    lt.thu_tu_dung
FROM LichTrinhChuyen lt
JOIN GaTau g ON lt.id_ga = g.id_ga
WHERE lt.id_lich_chay = @idLichChay
  AND lt.id_ga IN (@idGaLen, @idGaXuong);

-- Nếu query trên trả về < 2 dòng → giá sẽ null (lỗi nghiêm trọng)
-- Cách tính km: km = |km_ga_xuong - km_ga_len|

-- 2.1.3 Kiểm tra xem có đủ 2 điểm để tính km không
SELECT
    lc.id_lich_chay,
    t.so_hieu,
    COUNT(lt.id_lich_trinh) AS so_diem_dung,
    MIN(lt.khoang_cach_km) AS km_min,
    MAX(lt.khoang_cach_km) AS km_max
FROM LichChay lc
JOIN Tau t ON lc.id_tau = t.id_tau
LEFT JOIN LichTrinhChuyen lt ON lc.id_lich_chay = lt.id_lich_chay
GROUP BY lc.id_lich_chay, t.so_hieu
ORDER BY so_diem_dung;
-- Cảnh báo nếu so_diem_dung < 2""")

    # 2.2
    add_h2(doc, "2.2 Kiểm tra biểu giá (BieuGia)")
    add_code_block(doc, """\
-- 2.2.1 Xem tất cả biểu giá
SELECT
    id_bieu_gia,
    ten_bieu_gia,
    don_gia_km_goc,
    he_so_tang,
    ngay_bat_dau,
    ngay_ket_thuc,
    trang_thai,
    id_loai_ghe
FROM BieuGia
ORDER BY ngay_bat_dau DESC;

-- 2.2.2 Kiểm tra biểu giá đang áp dụng cho ngày hiện tại (QUAN TRỌNG)
DECLARE @ngayKiemTra DATE = CAST(GETDATE() AS DATE);
SELECT
    id_bieu_gia,
    ten_bieu_gia,
    don_gia_km_goc,
    he_so_tang,
    ngay_bat_dau,
    ngay_ket_thuc,
    trang_thai
FROM BieuGia
WHERE trang_thai = 'dang_ap_dung'
  AND ngay_bat_dau <= @ngayKiemTra
  AND ngay_ket_thuc >= @ngayKiemTra
ORDER BY he_so_tang DESC;
-- Nếu không có dòng nào → priceFrom luôn = null, backend dùng giá mặc định 264đ/km

-- 2.2.3 Mô phỏng tính giá vé (giống hàm tinhGiaVe của backend)
DECLARE @km          FLOAT = 1726;  -- km Hà Nội - Sài Gòn
DECLARE @donGia      FLOAT = 264;   -- Thay bằng don_gia_km_goc từ BieuGia
DECLARE @hesoTang    FLOAT = 1.0;   -- he_so_tang từ BieuGia
DECLARE @hesoGhe     FLOAT = 1.0;   -- he_so_ghe từ LoaiGhe

SELECT
    @km               AS km,
    @donGia           AS don_gia_km,
    @hesoTang         AS he_so_tang,
    @hesoGhe          AS he_so_ghe,
    CEILING((@km * @donGia * @hesoTang * @hesoGhe) / 1000) * 1000 AS gia_ve_tinh_duoc;""")

    # 2.3
    add_h2(doc, "2.3 Mô phỏng toàn bộ luồng tìm chuyến")
    add_code_block(doc, """\
-- 2.3.1 Toàn bộ query giống như backend thực hiện khi tìm chuyến
-- Thay tên ga và ngày phù hợp với dữ liệu của bạn
DECLARE @tenGaDi  NVARCHAR(100) = N'Hà Nội';
DECLARE @tenGaDen NVARCHAR(100) = N'Sài Gòn';
DECLARE @ngayChay DATE = '2026-06-15';

-- Bước 1: Tìm id ga
SELECT id_ga, ten_ga, trang_thai
FROM GaTau
WHERE ten_ga IN (@tenGaDi, @tenGaDen) AND trang_thai = 'hoat_dong';

-- Bước 2: Tìm lịch chạy (kết hợp bước 1)
SELECT
    ct.id_chuyen,
    ct.ngay_chay,
    t.so_hieu AS ma_tau,
    t.ten_tau,
    gd.ten_ga AS ga_di,
    gn.ten_ga AS ga_den,
    lc.gio_khoi_hanh,
    lc.gio_du_kien_den,
    ct.trang_thai
FROM ChuyenTau ct
JOIN LichChay lc ON ct.id_lich_chay = lc.id_lich_chay
JOIN Tau t ON lc.id_tau = t.id_tau
JOIN GaTau gd ON lc.id_ga_di = gd.id_ga
JOIN GaTau gn ON lc.id_ga_den = gn.id_ga
WHERE ct.ngay_chay = @ngayChay
  AND ct.trang_thai != 'huy'
  AND gd.ten_ga = @tenGaDi
  AND gn.ten_ga = @tenGaDen
ORDER BY lc.gio_khoi_hanh;

-- 2.3.2 Kiểm tra ghế trống của một chuyến (kết hợp CauHinhGhe + Ve + TamGiuGhe)
DECLARE @idChuyen    INT = 1;  -- Thay bằng id_chuyen thực tế
DECLARE @soToa       INT = 1;  -- Số thứ tự toa

-- Ghế đã đặt
SELECT so_ghe_trong_toa, trang_thai
FROM Ve
WHERE id_chuyen = @idChuyen
  AND so_toa_thu_tu = @soToa
  AND trang_thai NOT IN ('da_huy', 'da_doi');

-- Ghế đang giữ tạm (còn hiệu lực)
SELECT so_ghe_trong_toa, thoi_gian_het_han, trang_thai
FROM TamGiuGhe
WHERE id_chuyen = @idChuyen
  AND so_toa_thu_tu = @soToa
  AND trang_thai = 'dang_giu'
  AND thoi_gian_het_han > GETDATE();

-- Tổng hợp trạng thái từng ghế của toa
SELECT
    cg.so_ghe_trong_toa,
    lg.ten_loai_ghe,
    lg.he_so_ghe,
    cg.khoang_so,
    cg.ben,
    CASE
        WHEN v.so_ghe_trong_toa IS NOT NULL THEN 'sold'
        WHEN tg.so_ghe_trong_toa IS NOT NULL THEN 'held'
        ELSE 'empty'
    END AS trang_thai_ghe
FROM CauHinhGhe cg
JOIN LoaiGhe lg ON cg.id_loai_ghe = lg.id_loai_ghe
JOIN CauHinhToa ct ON cg.id_loai_toa = ct.id_loai_toa
JOIN ChuyenTau chu ON ct.id_tau = (
    SELECT lc2.id_tau FROM LichChay lc2
    JOIN ChuyenTau ct2 ON lc2.id_lich_chay = ct2.id_lich_chay
    WHERE ct2.id_chuyen = @idChuyen
)
AND ct.so_toa_thu_tu = @soToa
LEFT JOIN Ve v ON v.id_chuyen = @idChuyen
    AND v.so_toa_thu_tu = @soToa
    AND v.so_ghe_trong_toa = cg.so_ghe_trong_toa
    AND v.trang_thai NOT IN ('da_huy','da_doi')
LEFT JOIN TamGiuGhe tg ON tg.id_chuyen = @idChuyen
    AND tg.so_toa_thu_tu = @soToa
    AND tg.so_ghe_trong_toa = cg.so_ghe_trong_toa
    AND tg.trang_thai = 'dang_giu'
    AND tg.thoi_gian_het_han > GETDATE()
ORDER BY cg.so_ghe_trong_toa;""")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # PHẦN 3
    # ═══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "PHẦN 3: KIỂM TRA CHỨC NĂNG ĐẶT VÉ")

    add_h2(doc, "3.1 Kiểm tra đơn đặt vé")
    add_code_block(doc, """\
-- 3.1.1 Xem tất cả đơn đặt vé gần đây
SELECT TOP 20
    d.id_don_dat_ve,
    d.ma_dat_cho,
    d.ma_don,
    d.ho_ten_lien_lac,
    d.email_dat_cho,
    d.sdt_dat_cho,
    d.loai_ve,
    d.tong_tien,
    d.tien_giam,
    d.tien_thanh_toan,
    d.trang_thai,
    d.thoi_gian_dat,
    d.thoi_gian_het_han
FROM DonDatVe d
ORDER BY d.thoi_gian_dat DESC;

-- 3.1.2 Xem chi tiết đơn + vé (tra cứu như API lookupBooking)
DECLARE @maDatCho NVARCHAR(10) = 'ABC123'; -- Thay bằng mã thực tế
SELECT
    d.ma_dat_cho,
    d.ma_don,
    d.ho_ten_lien_lac,
    d.email_dat_cho,
    d.sdt_dat_cho,
    d.tong_tien,
    d.tien_thanh_toan,
    d.trang_thai AS trang_thai_don,
    v.id_ve,
    hk.ho_ten AS ten_hanh_khach,
    hk.ngay_sinh,
    hk.cccd,
    v.so_toa_thu_tu,
    v.so_ghe_trong_toa,
    v.gia_ve,
    v.trang_thai AS trang_thai_ve,
    chu.ngay_chay,
    lc.gio_khoi_hanh,
    lc.gio_du_kien_den,
    t.so_hieu AS ma_tau,
    gd.ten_ga AS ga_di,
    gn.ten_ga AS ga_den
FROM DonDatVe d
LEFT JOIN Ve v ON d.id_don_dat_ve = v.id_don_dat_ve
LEFT JOIN HanhKhach hk ON v.id_hanh_khach = hk.id_hanh_khach
LEFT JOIN ChuyenTau chu ON v.id_chuyen = chu.id_chuyen
LEFT JOIN LichChay lc ON chu.id_lich_chay = lc.id_lich_chay
LEFT JOIN Tau t ON lc.id_tau = t.id_tau
LEFT JOIN GaTau gd ON lc.id_ga_di = gd.id_ga
LEFT JOIN GaTau gn ON lc.id_ga_den = gn.id_ga
WHERE d.ma_dat_cho = @maDatCho;

-- 3.1.3 Kiểm tra trạng thái thanh toán của đơn
DECLARE @idDon INT = 1;
SELECT
    d.ma_dat_cho,
    d.trang_thai AS trang_thai_don,
    d.tien_thanh_toan,
    tt.id_thanh_toan,
    tt.phuong_thuc,
    tt.so_tien,
    tt.trang_thai AS trang_thai_thanh_toan,
    tt.thoi_gian_thanh_toan,
    hd.so_hoa_don
FROM DonDatVe d
LEFT JOIN ThanhToan tt ON d.id_don_dat_ve = tt.id_don_dat_ve
LEFT JOIN HoaDon hd ON tt.id_thanh_toan = hd.id_thanh_toan
WHERE d.id_don_dat_ve = @idDon;""")

    add_h2(doc, "3.2 Kiểm tra giữ ghế tạm thời (TamGiuGhe)")
    add_code_block(doc, """\
-- 3.2.1 Xem ghế đang bị giữ (chưa hết hạn)
SELECT
    tg.id_tam_giu,
    tg.id_chuyen,
    tg.so_toa_thu_tu,
    tg.so_ghe_trong_toa,
    tg.ma_dat_cho,
    tg.trang_thai,
    tg.thoi_gian_het_han,
    DATEDIFF(MINUTE, GETDATE(), tg.thoi_gian_het_han) AS phut_con_lai
FROM TamGiuGhe tg
WHERE tg.trang_thai = 'dang_giu'
  AND tg.thoi_gian_het_han > GETDATE()
ORDER BY tg.thoi_gian_het_han;

-- 3.2.2 Dọn dẹp ghế giữ đã hết hạn (nên chạy định kỳ hoặc khi cần debug)
UPDATE TamGiuGhe
SET trang_thai = 'da_giai_phong'
WHERE trang_thai = 'dang_giu'
  AND thoi_gian_het_han < GETDATE();
SELECT @@ROWCOUNT AS so_ghe_da_giai_phong;""")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # PHẦN 4
    # ═══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "PHẦN 4: KIỂM TRA CHỨC NĂNG TRA CỨU ĐẶT CHỖ")

    add_h2(doc, "4.1 Tra cứu đơn theo mã đặt chỗ + email + SĐT")
    add_code_block(doc, """\
-- 4.1.1 Mô phỏng API lookupBooking (phải khớp cả 3 điều kiện)
DECLARE @ma    NVARCHAR(10)  = 'ABC123';
DECLARE @email NVARCHAR(100) = 'khach@email.com';
DECLARE @sdt   NVARCHAR(15)  = '0901234567';

SELECT
    d.id_don_dat_ve,
    d.ma_dat_cho,
    d.ma_don,
    d.ho_ten_lien_lac,
    d.email_dat_cho,
    d.sdt_dat_cho,
    d.trang_thai,
    d.thoi_gian_dat
FROM DonDatVe d
WHERE d.ma_dat_cho = @ma
  AND d.email_dat_cho = @email
  AND REPLACE(d.sdt_dat_cho, ' ', '') = REPLACE(@sdt, ' ', '');
-- Nếu không ra kết quả: kiểm tra xem sdt có khoảng trắng hay dấu cách không

-- 4.1.2 Debug: tra cứu rộng hơn để tìm đơn
SELECT ma_dat_cho, ma_don, email_dat_cho, sdt_dat_cho, trang_thai
FROM DonDatVe
WHERE ma_dat_cho = @ma;

-- 4.1.3 Kiểm tra lịch sử đặt vé theo email (cho chức năng BookingHistory)
DECLARE @emailUser NVARCHAR(100) = 'khach@email.com';
SELECT
    d.ma_dat_cho,
    d.ma_don,
    d.tong_tien,
    d.tien_thanh_toan,
    d.trang_thai,
    d.thoi_gian_dat,
    COUNT(v.id_ve) AS so_ve
FROM DonDatVe d
LEFT JOIN Ve v ON d.id_don_dat_ve = v.id_don_dat_ve
WHERE d.email_dat_cho = @emailUser
GROUP BY d.id_don_dat_ve, d.ma_dat_cho, d.ma_don,
         d.tong_tien, d.tien_thanh_toan, d.trang_thai, d.thoi_gian_dat
ORDER BY d.thoi_gian_dat DESC;""")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # PHẦN 5
    # ═══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "PHẦN 5: KIỂM TRA CHỨC NĂNG HỦY VÉ")

    add_h2(doc, "5.1 Kiểm tra chính sách hủy (ChinhSachHuy)")
    add_code_block(doc, """\
-- 5.1.1 Xem toàn bộ chính sách hủy vé
SELECT
    id_cs_huy,
    gio_truoc_gio_chay,
    phi_huy,
    mo_ta
FROM ChinhSachHuy
ORDER BY gio_truoc_gio_chay;
-- Kết quả mong đợi: 0h=100%, 4h=50%, 24h=25%, 72h=10%

-- 5.1.2 Nếu bảng ChinhSachHuy rỗng → hủy vé sẽ luôn tính phí 100%
-- Chèn dữ liệu mẫu:
/*
IF NOT EXISTS (SELECT 1 FROM ChinhSachHuy)
BEGIN
    INSERT INTO ChinhSachHuy (gio_truoc_gio_chay, phi_huy, mo_ta) VALUES
    (0,  100, N'Dưới 4 giờ hoặc sau khởi hành — không hoàn tiền'),
    (4,  50,  N'Trước 4 giờ đến dưới 1 ngày — hoàn 50%'),
    (24, 25,  N'Trước 1 ngày đến dưới 3 ngày — hoàn 75%'),
    (72, 10,  N'Trước 3 ngày trở lên — hoàn 90%');
END
*/

-- 5.1.3 Mô phỏng tra cứu chính sách hủy (giống hàm tinhPhiHuy)
DECLARE @idVeHuy INT = 1;  -- Thay bằng id_ve cần hủy
DECLARE @gioConLai FLOAT;

SELECT
    v.id_ve,
    v.gia_ve,
    v.trang_thai,
    chu.ngay_chay,
    lc.gio_khoi_hanh,
    DATEDIFF(HOUR, GETDATE(),
        CAST(CAST(chu.ngay_chay AS VARCHAR) + ' ' + CAST(lc.gio_khoi_hanh AS VARCHAR) AS DATETIME)
    ) AS gio_con_lai
FROM Ve v
JOIN ChuyenTau chu ON v.id_chuyen = chu.id_chuyen
JOIN LichChay lc ON chu.id_lich_chay = lc.id_lich_chay
WHERE v.id_ve = @idVeHuy;

-- 5.1.4 Tính phí hủy theo giờ còn lại
DECLARE @gioTruocChoChay INT = 30; -- Thay bằng gio_con_lai thực tế
SELECT TOP 1
    gio_truoc_gio_chay,
    phi_huy,
    mo_ta,
    (100 - phi_huy) AS phan_tram_hoan
FROM ChinhSachHuy
WHERE gio_truoc_gio_chay <= @gioTruocChoChay
ORDER BY gio_truoc_gio_chay DESC;
-- Nếu query trên rỗng → áp dụng phí 100% (không hoàn)

-- 5.1.5 Xem vé và thông tin hoàn tiền sau khi hủy
SELECT
    v.id_ve,
    v.gia_ve,
    v.trang_thai,
    ht.tien_goc,
    ht.phi_huy,
    ht.tien_hoan,
    ht.trang_thai_hoan,
    ht.thoi_gian_hoan
FROM Ve v
LEFT JOIN HoanTien ht ON v.id_ve = ht.id_ve
WHERE v.trang_thai = 'da_huy'
ORDER BY ht.thoi_gian_hoan DESC;""")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # PHẦN 6
    # ═══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "PHẦN 6: KIỂM TRA CHỨC NĂNG ĐỔI VÉ")

    add_h2(doc, "6.1 Kiểm tra điều kiện và kết quả đổi vé")
    add_code_block(doc, """\
-- 6.1.1 Mô phỏng checkExchangeable: kiểm tra vé có thể đổi không
DECLARE @idVeDoi INT = 1;  -- Thay bằng id_ve cần đổi
SELECT
    v.id_ve,
    v.gia_ve,
    v.trang_thai,
    chu.ngay_chay,
    lc.gio_khoi_hanh,
    DATEDIFF(HOUR, GETDATE(),
        CAST(CAST(chu.ngay_chay AS VARCHAR) + ' ' + CAST(lc.gio_khoi_hanh AS VARCHAR) AS DATETIME)
    ) AS gio_con_lai,
    CASE
        WHEN v.trang_thai != 'da_xac_nhan' THEN N'Không đổi được: vé chưa xác nhận'
        WHEN DATEDIFF(HOUR, GETDATE(),
            CAST(CAST(chu.ngay_chay AS VARCHAR) + ' ' + CAST(lc.gio_khoi_hanh AS VARCHAR) AS DATETIME)
        ) < 24 THEN N'Không đổi được: dưới 24 giờ khởi hành'
        ELSE N'Có thể đổi vé'
    END AS ket_qua_kiem_tra
FROM Ve v
JOIN ChuyenTau chu ON v.id_chuyen = chu.id_chuyen
JOIN LichChay lc ON chu.id_lich_chay = lc.id_lich_chay
WHERE v.id_ve = @idVeDoi;

-- 6.1.2 Xem lịch sử đổi vé
SELECT
    dv.id_doi_ve,
    v_cu.id_ve AS id_ve_cu,
    v_moi.id_ve AS id_ve_moi,
    dv.phi_doi,
    dv.chenh_lech_gia,
    dv.tong_phai_tra,
    dv.trang_thai,
    dv.thoi_gian_doi
FROM DoiVe dv
JOIN Ve v_cu ON dv.id_ve_cu = v_cu.id_ve
JOIN Ve v_moi ON dv.id_ve_moi = v_moi.id_ve
ORDER BY dv.thoi_gian_doi DESC;

-- 6.1.3 Tính phí đổi vé (5%, tối thiểu 20.000đ)
DECLARE @giaVeCu FLOAT = 456000;
SELECT
    @giaVeCu AS gia_ve_cu,
    CASE
        WHEN @giaVeCu * 0.05 < 20000 THEN 20000
        ELSE ROUND(@giaVeCu * 0.05, 0)
    END AS phi_doi_ve;""")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # PHẦN 7
    # ═══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "PHẦN 7: KIỂM TRA KHUYẾN MÃI")

    add_code_block(doc, """\
-- 7.1 Xem tất cả mã khuyến mãi
SELECT
    id_khuyen_mai,
    ma_khuyen_mai,
    ten_khuyen_mai,
    loai_giam,
    gia_tri,
    giam_toi_da,
    gia_tri_don_toi_thieu,
    ngay_bat_dau,
    ngay_het_han,
    so_luong,
    da_dung,
    trang_thai
FROM KhuyenMai
ORDER BY ngay_het_han DESC;

-- 7.2 Kiểm tra mã khuyến mãi còn hiệu lực
DECLARE @ngayHienTai DATE = CAST(GETDATE() AS DATE);
SELECT *
FROM KhuyenMai
WHERE trang_thai = 'hoat_dong'
  AND ngay_bat_dau <= @ngayHienTai
  AND ngay_het_han >= @ngayHienTai
  AND (so_luong IS NULL OR da_dung < so_luong);

-- 7.3 Tính giá sau khi áp khuyến mãi
DECLARE @tongDon   FLOAT = 1000000;
DECLARE @loaiGiam  VARCHAR(20) = 'phan_tram';
DECLARE @giaTriKM  FLOAT = 10;      -- 10%
DECLARE @giamToiDa FLOAT = 50000;   -- Tối đa 50k

SELECT
    @tongDon AS tong_don,
    CASE
        WHEN @loaiGiam = 'phan_tram' THEN
            LEAST(FLOOR(@tongDon * @giaTriKM / 100), @giamToiDa)
        ELSE @giaTriKM
    END AS tien_giam,
    @tongDon - CASE
        WHEN @loaiGiam = 'phan_tram' THEN
            LEAST(FLOOR(@tongDon * @giaTriKM / 100), @giamToiDa)
        ELSE @giaTriKM
    END AS tong_sau_giam;""")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # PHẦN 8
    # ═══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "PHẦN 8: CHÈN DỮ LIỆU MẪU ĐẦY ĐỦ")
    add_normal(doc,
        "Chạy phần này nếu CSDL còn thiếu dữ liệu để test. "
        "Kiểm tra từng bước trước khi chạy.")

    add_h2(doc, "8.1 Dữ liệu mẫu: BieuGia")
    add_code_block(doc, """\
-- Kiểm tra trước
SELECT COUNT(*) AS so_bieu_gia FROM BieuGia WHERE trang_thai = 'dang_ap_dung';

-- Nếu = 0, chèn biểu giá mặc định
/*
INSERT INTO BieuGia (ten_bieu_gia, don_gia_km_goc, he_so_tang,
                     ngay_bat_dau, ngay_ket_thuc, trang_thai, id_loai_ghe)
VALUES
(N'Biểu giá thường 2026', 264, 1.0, '2026-01-01', '2026-12-31', 'dang_ap_dung', NULL),
(N'Biểu giá hè 2026',     264, 1.2, '2026-06-01', '2026-08-31', 'dang_ap_dung', NULL);
*/""")

    add_h2(doc, "8.2 Dữ liệu mẫu: ChinhSachHuy")
    add_code_block(doc, """\
-- Kiểm tra trước
SELECT COUNT(*) AS so_chinh_sach FROM ChinhSachHuy;

-- Nếu = 0, chèn chính sách hủy
/*
INSERT INTO ChinhSachHuy (gio_truoc_gio_chay, phi_huy, mo_ta)
VALUES
(0,  100, N'Dưới 4 giờ hoặc sau khởi hành — không hoàn tiền'),
(4,  50,  N'Trước 4 giờ đến 1 ngày — hoàn 50%'),
(24, 25,  N'Trước 1 ngày đến 3 ngày — hoàn 75%'),
(72, 10,  N'Trước 3 ngày trở lên — hoàn 90%');
*/""")

    add_h2(doc, "8.3 Dữ liệu mẫu: ChuyenTau cho các ngày test")
    add_code_block(doc, """\
-- Kiểm tra chuyến trong tháng tới
DECLARE @tuNgay DATE = CAST(GETDATE() AS DATE);
DECLARE @denNgay DATE = DATEADD(MONTH, 1, @tuNgay);

SELECT COUNT(*) AS so_chuyen_thang_toi
FROM ChuyenTau
WHERE ngay_chay BETWEEN @tuNgay AND @denNgay
  AND trang_thai != 'huy';

-- Nếu quá ít chuyến, tạo thêm từ lịch chạy hiện có
/*
-- Tạo chuyến cho từng ngày trong 30 ngày tới (chạy với từng id_lich_chay)
DECLARE @idLichChay INT = 1; -- Thay bằng id thực tế
DECLARE @i INT = 0;
WHILE @i < 30
BEGIN
    DECLARE @ngay DATE = DATEADD(DAY, @i, CAST(GETDATE() AS DATE));
    IF NOT EXISTS (SELECT 1 FROM ChuyenTau WHERE id_lich_chay = @idLichChay AND ngay_chay = @ngay)
    BEGIN
        INSERT INTO ChuyenTau (id_lich_chay, ngay_chay, trang_thai)
        VALUES (@idLichChay, @ngay, 'dang_chay');
    END
    SET @i = @i + 1;
END
*/""")

    add_h2(doc, "8.4 Kiểm tra LichTrinhChuyen và thêm nếu thiếu")
    add_code_block(doc, """\
-- Xem lịch trình hiện tại
SELECT
    lc.id_lich_chay,
    t.so_hieu,
    COUNT(lt.id_lich_trinh) AS so_diem_dung
FROM LichChay lc
JOIN Tau t ON lc.id_tau = t.id_tau
LEFT JOIN LichTrinhChuyen lt ON lc.id_lich_chay = lt.id_lich_chay
GROUP BY lc.id_lich_chay, t.so_hieu
ORDER BY so_diem_dung;

-- Nếu thiếu (< 2 điểm dừng), thêm điểm ga đi (km=0) và ga đến
-- Ví dụ: Hà Nội (km=0) → Sài Gòn (km=1726)
/*
DECLARE @idLichChay INT = 1;
DECLARE @idGaHN INT = 1;  -- id_ga của Hà Nội
DECLARE @idGaSG INT = 2;  -- id_ga của Sài Gòn

INSERT INTO LichTrinhChuyen (id_lich_chay, id_ga, thu_tu_dung, khoang_cach_km,
                              gio_du_kien_den, gio_du_kien_di)
VALUES
(@idLichChay, @idGaHN, 1, 0,    NULL,    '06:00:00'),
(@idLichChay, @idGaSG, 2, 1726, '04:00:00', NULL);
*/""")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # PHẦN 9
    # ═══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "PHẦN 9: BÁO CÁO TỔNG QUAN HỆ THỐNG")

    add_code_block(doc, """\
-- 9.1 Thống kê tổng quan (dashboard kiểm tra sức khỏe DB)
SELECT
    (SELECT COUNT(*) FROM GaTau WHERE trang_thai = 'hoat_dong') AS so_ga,
    (SELECT COUNT(*) FROM Tau) AS so_tau,
    (SELECT COUNT(*) FROM LichChay WHERE trang_thai != 'huy') AS so_lich_chay,
    (SELECT COUNT(*) FROM ChuyenTau WHERE ngay_chay >= CAST(GETDATE() AS DATE) AND trang_thai != 'huy') AS chuyen_tuong_lai,
    (SELECT COUNT(*) FROM DonDatVe) AS tong_don,
    (SELECT COUNT(*) FROM DonDatVe WHERE trang_thai = 'cho_thanh_toan') AS don_cho_tt,
    (SELECT COUNT(*) FROM DonDatVe WHERE trang_thai IN ('da_thanh_toan','da_xac_nhan')) AS don_da_tt,
    (SELECT COUNT(*) FROM Ve WHERE trang_thai = 'da_huy') AS ve_da_huy,
    (SELECT COUNT(*) FROM BieuGia WHERE trang_thai = 'dang_ap_dung') AS bieu_gia_active,
    (SELECT COUNT(*) FROM ChinhSachHuy) AS chinh_sach_huy,
    (SELECT COUNT(*) FROM TamGiuGhe WHERE trang_thai = 'dang_giu' AND thoi_gian_het_han > GETDATE()) AS ghe_dang_giu;

-- 9.2 Doanh thu theo ngày
SELECT
    CAST(tt.thoi_gian_thanh_toan AS DATE) AS ngay,
    COUNT(*) AS so_giao_dich,
    SUM(tt.so_tien) AS tong_doanh_thu
FROM ThanhToan tt
WHERE tt.trang_thai = 'thanh_cong'
GROUP BY CAST(tt.thoi_gian_thanh_toan AS DATE)
ORDER BY ngay DESC;

-- 9.3 Phát hiện dữ liệu bất thường
-- Vé không có đơn đặt (orphaned records)
SELECT v.id_ve, v.id_don_dat_ve
FROM Ve v
WHERE NOT EXISTS (SELECT 1 FROM DonDatVe d WHERE d.id_don_dat_ve = v.id_don_dat_ve);

-- Đơn quá hạn thanh toán vẫn còn trạng thái cho_thanh_toan
SELECT id_don_dat_ve, ma_dat_cho, thoi_gian_het_han, trang_thai
FROM DonDatVe
WHERE trang_thai = 'cho_thanh_toan'
  AND thoi_gian_het_han < GETDATE()
ORDER BY thoi_gian_het_han;

-- 9.4 Cập nhật đơn hết hạn (nên chạy định kỳ)
/*
UPDATE DonDatVe
SET trang_thai = 'het_han'
WHERE trang_thai = 'cho_thanh_toan'
  AND thoi_gian_het_han < GETDATE();
SELECT @@ROWCOUNT AS so_don_da_cap_nhat;
*/""")

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run("— Hết tài liệu — KLN Train SQL Test Queries © 2026 —")
    set_run_font(r, 'Calibri', 10, color_rgb=(0x80, 0x80, 0x80))

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"[OK] Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
